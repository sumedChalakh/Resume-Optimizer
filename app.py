import os
import json
import re
import urllib.request
import urllib.error
from datetime import date
from collections import Counter
from flask import Flask, render_template, request, jsonify, send_file
from anthropic import Anthropic
from docx import Document
from pypdf import PdfReader
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
import io

from extensions import db, bcrypt, login_manager
from models import User

def load_env_file(path=".env"):
  """Load KEY=VALUE pairs from a local .env file into process environment."""
  if not os.path.exists(path):
    return

  with open(path, "r", encoding="utf-8") as env_file:
    for raw_line in env_file:
      line = raw_line.strip()
      if not line or line.startswith("#") or "=" not in line:
        continue
      key, value = line.split("=", 1)
      key = key.strip()
      value = value.strip().strip('"').strip("'")
      if key and key not in os.environ:
        os.environ[key] = value

load_env_file()



# session and admin dashboard endpoints moved to auth.routes


# `login_page` moved to `auth.routes` blueprint




try:
  MAX_UPLOAD_BYTES = int(os.getenv("MAX_UPLOAD_BYTES", str(8 * 1024 * 1024)))
except ValueError:
  MAX_UPLOAD_BYTES = 8 * 1024 * 1024

ANTHROPIC_API_KEY = os.getenv("ANTHROPIC_API_KEY", "").strip()
GITHUB_PAT = os.getenv("GITHUB_PAT", "").strip()
OPENROUTER_API_KEY = os.getenv("OPENROUTER_API_KEY", "").strip()
OPENROUTER_API_KEY_BACKUP = os.getenv("OPENROUTER_API_KEY_BACKUP", "").strip()
HUGGINGFACE_API_KEY = os.getenv("HUGGINGFACE_API_KEY", "").strip()
HUGGINGFACE_API_KEY_BACKUP = os.getenv("HUGGINGFACE_API_KEY_BACKUP", "").strip()
GITHUB_MODEL = os.getenv("GITHUB_MODEL", "gpt-4o-mini").strip()
HUGGINGFACE_MODEL = os.getenv("HUGGINGFACE_MODEL", "meta-llama/Llama-3.1-8B-Instruct").strip()
OPENROUTER_MODEL = os.getenv("OPENROUTER_MODEL", "deepseek/deepseek-r1:free").strip()
MODEL_TIMEOUT_SECONDS = 75
MODEL_MAX_TOKENS = 2600
MODEL_INPUT_MAX_CHARS = 8500
OPENROUTER_LOW_BALANCE_MAX_TOKENS = 850
OPENROUTER_MODEL_FALLBACKS = [
  "meta-llama/llama-3.3-70b-instruct:free",
  "openrouter/auto",
]
HUGGINGFACE_MAX_NEW_TOKENS = 1100
HUGGINGFACE_TIMEOUT_SECONDS = 90
HUGGINGFACE_CHAT_FALLBACK_MODELS = [
  "meta-llama/Llama-3.1-8B-Instruct",
  "Qwen/Qwen2.5-7B-Instruct",
]
GITHUB_MODELS_FALLBACKS = [
  "gpt-4.1-mini",
  "gpt-4o-mini",
]

client = (
  Anthropic(api_key=ANTHROPIC_API_KEY, timeout=MODEL_TIMEOUT_SECONDS, max_retries=0)
  if ANTHROPIC_API_KEY
  else None
)

SYSTEM_PROMPT = """You are a world-class ATS Resume Optimizer, Technical Recruiter, and Career Coach.
Your task is to optimize a candidate resume for a target job description while staying truthful, ATS-safe, and recruiter-friendly.
CORE RULES:
1. Truth first: never fabricate experience, achievements, or certifications.
2. Keyword coverage: extract and naturally use important JD keywords across summary, skills, experience, and projects.
3. Preserve entries: NEVER remove original projects or certifications. Keep every original project and certification entry; rewrite and reorder for JD relevance only.
4. Skill completeness:
   - Keep all original resume skills.
   - Add skills clearly inferred from projects/experience.
   - Add strong JD-aligned adjacent skills only when realistic, and explain those additions.
5. Bullet quality: use strong action verbs, concrete technologies, and measurable impact where possible.
6. ATS format: use standard sections and plain ATS-readable wording.

OUTPUT REQUIREMENTS:
-- Respond ONLY with valid JSON (no markdown or extra text).
-- Keep the schema below exactly.

{
  "optimized_resume": {
    "summary": "3-line recruiter-optimized summary",
    "skills": {
      "Programming Languages": ["Python", "SQL"],
      "Frameworks & Libraries": ["scikit-learn", "TensorFlow"],
      "Machine Learning & AI": ["Supervised Learning", "XGBoost"],
      "Data Engineering & Big Data": ["PySpark", "Kafka"],
      "Tools & Platforms": ["Docker", "Git", "Jupyter"],
      "Databases": ["MySQL", "PostgreSQL"],
      "Cloud & DevOps": ["AWS", "CI/CD"],
      "Data Visualization": ["Power BI", "Matplotlib"],
      "Soft Skills": ["Problem Solving", "Communication"]
    },
    "experience": [
      {
        "title": "Job Title",
        "company": "Company Name",
        "duration": "Month Year - Month Year",
        "bullets": [
          "Action + technology + quantified impact",
          "Action + technical detail + business outcome",
          "Action + optimization + measurable result"
        ]
      }
    ],
    "projects": [
      {
        "name": "Project Name",
        "tech": "Comma-separated full stack",
        "bullets": [
          "Action + build detail + scale/metric",
          "Action + technical depth + impact"
        ]
      }
    ],
    "education": [
      {
        "degree": "Full Degree Name",
        "institution": "Institution Name",
        "year": "Year",
        "details": "Relevant coursework or honors"
      }
    ],
    "certifications": [
      {
        "name": "Certification Name",
        "issuer": "Organization",
        "year": "Year"
      }
    ]
  },
  "ats_score": {
    "total": 87,
    "breakdown": {
      "keyword_match": 25,
      "skills_completeness": 18,
      "bullet_quality": 13,
      "summary_relevance": 13,
      "structure_parsability": 9,
      "education_certifications": 9
    },
    "score_reasoning": "Short explanation"
  },
  "keyword_analysis": {
    "jd_keywords_extracted": ["kw1", "kw2", "kw3"],
    "matched_in_resume": ["kw1", "kw2"],
    "missing_keywords": ["kw3"],
    "keyword_match_percentage": 85
  },
  "skills_audit": {
    "skills_from_original_resume": ["skill1"],
    "skills_inferred_from_projects": ["skill2"],
    "skills_added_from_jd": ["skill3"],
    "skills_removed": [],
    "note": "All original skills retained"
  },
  "improvements": [
    "[KEYWORD] Specific keyword gap",
    "[SKILL ADD] Skill added from JD",
    "[SKILL INFERRED] Skill inferred from evidence",
    "[BULLET] Bullet improvement",
    "[SUMMARY] Summary alignment fix",
    "[STRUCTURE] ATS structure optimization"
  ],
  "recruiter_tips": [
    "Interview positioning tip",
    "Best project to lead with",
    "Optional cover letter angle"
  ],
  "analysis": {
    "role_fit_score": "Strong",
    "top_strengths": ["Strength 1", "Strength 2"],
    "weak_points": ["Gap 1", "Gap 2"],
    "differentiators": ["Unique differentiator"]
  }
}"""


COVER_LETTER_PROMPT = """You are an expert cover letter writer.
Write a professional, ATS-friendly cover letter based on the resume and job description provided.

STRICT RULES:
-- Maximum 4 paragraphs
-- Para 1: Hook — why THIS company + role excites you (mention company name)
-- Para 2: Your strongest 2-3 achievements relevant to THIS JD
-- Para 3: What unique value you bring (domain + tech stack match)
-- Para 4: Call to action — confident closing
-- IMPORTANT: DO NOT include the greeting ("Dear Hiring Manager,") or sign-off ("Sincerely, [Name]") in the text. Return ONLY the actual paragraphs.

TONE: Professional but human. Not robotic. Not over-formal.
LENGTH: 250-320 words max — recruiters don't read long letters

Respond ONLY in this JSON:
{
  "cover_letter": {
    "subject_line": "Application for [exact job title] — [Your Name]",
    "body_paragraphs": [
      "First paragraph here...",
      "Second paragraph here...",
      "Third paragraph here...",
      "Fourth paragraph here..."
    ],
    "word_count": 285
  },
  "personalization_score": 87,
  "tips": [
    "tip 1 to make it stronger",
  ]
}"""


ALIGNMENT_PROMPT = """You are an expert technical recruiter and ATS systems evaluator.
Perform a high-speed domain alignment check between the job title/description and the candidate's professional summary.
Your goal is to determine if the overall professional profile is heavily misaligned with the target job domain (e.g., applying for a Data Science role with a pure Nursing background).

Respond ONLY in this JSON format:
{
  "misaligned": true_or_false,
  "penalty": integer_between_0_and_30,
  "reasoning": "Brief explanation of the penalty."
}"""


def extract_certifications_from_resume_text(resume_text):
  """Best-effort extraction of certification lines from original resume text."""
  if not resume_text:
    return []

  lines = [line.strip(" -•\t") for line in str(resume_text).splitlines() if line.strip()]
  certs = []
  in_cert_section = False

  for line in lines:
    low = line.lower()
    if low in {"certifications", "certification", "licenses & certifications", "licenses and certifications"}:
      in_cert_section = True
      continue

    if in_cert_section and low in {"education", "projects", "experience", "skills", "summary", "professional summary"}:
      break

    if in_cert_section or any(k in low for k in ["certification", "certified", "course", "nptel", "aws", "azure", "google ai"]):
      if len(line) < 5:
        continue
      certs.append(line)

  cleaned = []
  seen = set()
  for cert in certs:
    key = cert.lower()
    if key in seen:
      continue
    seen.add(key)
    cleaned.append({"name": cert, "issuer": "", "year": ""})
  return cleaned


def canonicalize_cert_key(name):
  value = normalize_space(name).lower()
  if not value:
    return ""

  value = value.replace("–", "-").replace("—", "-")
  value = re.sub(r"\([^)]*\)", "", value)
  value = re.sub(r"^[a-z][a-z\s&/\+\-]{1,20}:\s*", "", value)
  value = re.sub(
    r"\b(?:jan|feb|mar|apr|may|jun|jul|aug|sep|sept|oct|nov|dec)\s+\d{4}\b",
    "",
    value,
    flags=re.IGNORECASE,
  )
  value = re.sub(r"\b\d{4}\b", "", value)
  parts = [p.strip() for p in re.split(r"\s+-\s+", value) if p.strip()]
  if parts:
    value = parts[0]
  value = re.sub(r"[^a-z0-9\+\s]", " ", value)
  value = re.sub(r"\s+", " ", value).strip()
  return value


def canonicalize_project_key(name):
  value = normalize_space(name).lower()
  if not value:
    return ""

  value = value.replace("–", "-").replace("—", "-")
  value = re.sub(
    r"\b(?:jan|feb|mar|apr|may|jun|jul|aug|sep|sept|oct|nov|dec)\s+\d{4}(?:\s*-\s*(?:present|(?:jan|feb|mar|apr|may|jun|jul|aug|sep|sept|oct|nov|dec)\s+\d{4}))?\b",
    "",
    value,
    flags=re.IGNORECASE,
  )
  value = re.sub(r"\b\d{4}\b", "", value)
  value = re.sub(r"\s+", " ", value).strip(" -")
  return value


def project_quality_score(project):
  bullets = project.get("bullets", []) if isinstance(project, dict) else []
  tech = normalize_space(project.get("tech", "")) if isinstance(project, dict) else ""
  name = normalize_space(project.get("name", "")) if isinstance(project, dict) else ""
  return len(bullets) * 3 + (2 if tech else 0) + (1 if len(name) > 20 else 0)


def normalize_for_compare(value):
  text = normalize_space(str(value)).lower()
  text = re.sub(r"[^a-z0-9\s]", " ", text)
  return re.sub(r"\s+", " ", text).strip()


def text_token_set(value):
  tokens = normalize_for_compare(value).split()
  return {t for t in tokens if len(t) > 2}


def is_near_duplicate_text(a, b):
  na = normalize_for_compare(a)
  nb = normalize_for_compare(b)
  if not na or not nb:
    return False
  if na == nb:
    return True
  if len(na) > 35 and na in nb:
    return True
  if len(nb) > 35 and nb in na:
    return True

  ta = text_token_set(na)
  tb = text_token_set(nb)
  if not ta or not tb:
    return False

  overlap = len(ta & tb)
  union = len(ta | tb)
  jaccard = overlap / union if union else 0
  return jaccard >= 0.82 and min(len(ta), len(tb)) >= 6


def clean_bullet_lines(lines, max_items=6):
  if not isinstance(lines, list):
    return []

  clean = []
  for raw in lines:
    text = normalize_space(str(raw))
    low = text.lower()
    if not text:
      continue
    if len(text) < 15:
      continue
    if low in {"projects", "project", "experience", "skills", "education", "certifications"}:
      continue
    if any(is_near_duplicate_text(text, existing) for existing in clean):
      continue
    clean.append(text)
    if len(clean) >= max_items:
      break

  return clean


def normalize_project_entry(project):
  if not isinstance(project, dict):
    return {"name": "", "tech": "", "bullets": []}

  bullets = clean_bullet_lines(project.get("bullets", []), max_items=6)

  return {
    "name": normalize_space(project.get("name", "")),
    "tech": normalize_space(project.get("tech", "")),
    "bullets": bullets,
  }


def looks_like_project_sentence(text):
  value = normalize_space(text)
  low = value.lower()
  if not value:
    return False

  action_prefixes = (
    "built", "developed", "implemented", "engineered", "optimized", "applied",
    "processed", "designed", "delivered", "trained", "secured", "containerized",
    "evaluated", "gained", "executed", "collaborated",
  )
  if len(value) > 95:
    return True
  if value.endswith(".") and len(value.split()) > 8:
    return True
  if low.startswith(action_prefixes) and len(value.split()) > 6:
    return True
  if value[:1].islower() and len(value.split()) > 5:
    return True
  if ";" in value and len(value.split()) > 6:
    return True
  if re.search(r"\b\d+(?:\.\d+)?%\b", value) and len(value.split()) > 6:
    return True
  return False


def is_valid_project_entry(project):
  if not isinstance(project, dict):
    return False

  name = normalize_space(project.get("name", ""))
  tech = normalize_space(project.get("tech", ""))
  bullets = project.get("bullets", [])
  if not isinstance(bullets, list):
    bullets = []

  if not name:
    return False
  if is_probable_date_line(name):
    return False
  if is_probable_tech_line(name):
    return False
  if looks_like_project_sentence(name):
    return False
  if len(name.split()) > 12:
    return False
  if name[:1].islower():
    return False
  if name.count("|") > 1:
    return False

  # Avoid including malformed entries that are likely bullet fragments.
  if (";" in name or "," in name) and len(name.split()) > 10:
    return False

  # Reject entries that are just stack lines unless they have meaningful bullets.
  if (not tech) and len(bullets) == 0 and len(name.split()) > 8:
    return False

  return True


def normalize_experience_entry(exp):
  if not isinstance(exp, dict):
    return {"title": "", "company": "", "duration": "", "bullets": []}

  return {
    "title": normalize_space(exp.get("title", "")),
    "company": normalize_space(exp.get("company", "")),
    "duration": normalize_space(exp.get("duration", "")),
    "bullets": clean_bullet_lines(exp.get("bullets", []), max_items=6),
  }


def sanitize_resume_entries(result):
  if not isinstance(result, dict):
    return result

  optimized = result.get("optimized_resume")
  if not isinstance(optimized, dict):
    return result

  projects = optimized.get("projects")
  if isinstance(projects, list):
    clean_projects = []
    seen_project_names = []
    for proj in projects:
      normalized = normalize_project_entry(proj)
      if not is_valid_project_entry(normalized):
        continue
      name = normalized.get("name", "")
      if any(is_near_duplicate_text(name, existing_name) for existing_name in seen_project_names):
        continue
      seen_project_names.append(name)
      clean_projects.append(normalized)
    optimized["projects"] = clean_projects

  experience = optimized.get("experience")
  if isinstance(experience, list):
    clean_experience = []
    for exp in experience:
      normalized = normalize_experience_entry(exp)
      if not normalized.get("title") and not normalized.get("company") and not normalized.get("bullets"):
        continue
      clean_experience.append(normalized)
    optimized["experience"] = clean_experience

  return result


def is_probable_tech_line(line):
  low = normalize_space(line).lower()
  if not low:
    return False

  hints = [
    "python", "sql", "pyspark", "spark", "docker", "power bi", "dax", "flask",
    "streamlit", "scikit", "tensorflow", "keras", "pytorch", "mllib", "pandas",
    "numpy", "matplotlib", "github actions", "rest api", "kubernetes", "aws",
  ]
  hits = sum(1 for hint in hints if hint in low)
  comma_count = low.count(",")
  return hits >= 2 or (hits >= 1 and comma_count >= 2)


def is_probable_date_line(line):
  low = normalize_space(line).lower()
  if not low:
    return False

  return bool(re.search(
    r"^(?:jan|feb|mar|apr|may|jun|jul|aug|sep|sept|oct|nov|dec)\s+\d{4}(?:\s*[\-–—]\s*(?:present|(?:jan|feb|mar|apr|may|jun|jul|aug|sep|sept|oct|nov|dec)\s+\d{4}))?$",
    low,
    flags=re.IGNORECASE,
  ))


def is_probable_project_subtitle(line):
  low = normalize_space(line).lower()
  if not low:
    return False
  if "," in low:
    return False
  return any(kw in low for kw in ["project", "pipeline", "dashboard", "api", "internship"])


def ensure_certifications_present(result, resume_text):
  if not isinstance(result, dict):
    return result
  optimized = result.get("optimized_resume")
  if not isinstance(optimized, dict):
    return result

  fallback = extract_certifications_from_resume_text(resume_text)
  existing = optimized.get("certifications")
  if not isinstance(existing, list):
    existing = []

  seen = set()
  merged = []

  for cert in existing:
    if isinstance(cert, str):
      name = normalize_space(cert)
      cert_obj = {"name": name, "issuer": "", "year": ""}
    elif isinstance(cert, dict):
      name = normalize_space(cert.get("name", ""))
      cert_obj = {
        "name": name,
        "issuer": normalize_space(cert.get("issuer", "")),
        "year": normalize_space(cert.get("year", "")),
      }
    else:
      continue

    key = canonicalize_cert_key(name)
    if not key or key in seen:
      continue
    seen.add(key)
    merged.append(cert_obj)

  for cert in fallback:
    name = normalize_space(cert.get("name", ""))
    key = canonicalize_cert_key(name)
    if not key or key in seen:
      continue
    seen.add(key)
    merged.append({
      "name": name,
      "issuer": normalize_space(cert.get("issuer", "")),
      "year": normalize_space(cert.get("year", "")),
    })

  optimized["certifications"] = merged
  return result


def extract_projects_from_resume_text(resume_text):
  """Best-effort extraction of project entries from original resume text."""
  if not resume_text:
    return []

  lines = [line.strip() for line in resume_text.splitlines() if line.strip()]
  projects = []
  current = None
  in_projects = False

  project_headers = {"projects", "project", "personal projects", "academic projects"}
  section_breaks = {
    "experience",
    "professional experience",
    "work experience",
    "education",
    "skills",
    "certifications",
    "certification",
    "achievements",
    "summary",
    "professional summary",
  }

  for raw in lines:
    line = raw.strip(" \t")
    low = line.lower().strip(" :")

    if low in project_headers:
      in_projects = True
      continue

    if in_projects and low in section_breaks:
      break

    if not in_projects:
      continue

    # Attach bullet lines to the active project entry.
    if line.startswith(("-", "•", "*", "▸")):
      if current is not None:
        bullet = normalize_space(line.lstrip("-•*▸ "))
        if bullet:
          current.setdefault("bullets", []).append(bullet)
      continue

    if current is not None and is_probable_tech_line(line):
      tech = normalize_space(current.get("tech", ""))
      incoming = normalize_space(line)
      current["tech"] = f"{tech}, {incoming}".strip(", ") if tech else incoming
      continue

    if current is not None and is_probable_date_line(line):
      continue

    if current is not None and is_probable_project_subtitle(line):
      continue

    candidate = normalize_space(line)
    if len(candidate) < 3:
      continue
    if any(tok in low for tok in ["linkedin", "github.com", "@"]):
      continue

    name = candidate
    tech = ""
    if "|" in candidate:
      left, right = candidate.split("|", 1)
      name = normalize_space(left)
      tech = normalize_space(right)

    current = {"name": name, "tech": tech, "bullets": []}
    projects.append(current)

  deduped = []
  seen = set()
  for proj in projects:
    key = canonicalize_project_key(proj.get("name", ""))
    if not key or key in seen:
      continue
    seen.add(key)
    deduped.append(normalize_project_entry(proj))

  return deduped


def ensure_projects_present(result, resume_text):
  if not isinstance(result, dict):
    return result
  optimized = result.get("optimized_resume")
  if not isinstance(optimized, dict):
    return result

  extracted_projects = extract_projects_from_resume_text(resume_text)
  existing = optimized.get("projects")
  if not isinstance(existing, list):
    existing = []

  by_key = {}
  order = []

  for proj in existing:
    if not isinstance(proj, dict):
      continue
    normalized = normalize_project_entry(proj)
    if not is_valid_project_entry(normalized):
      continue
    key = canonicalize_project_key(normalized.get("name", ""))
    if not key:
      continue
    if key not in by_key:
      by_key[key] = normalized
      order.append(key)
      continue

    current = by_key[key]
    if project_quality_score(normalized) > project_quality_score(current):
      replacement = normalized
      if not replacement.get("tech") and current.get("tech"):
        replacement["tech"] = current.get("tech")
      if not replacement.get("bullets") and current.get("bullets"):
        replacement["bullets"] = current.get("bullets")
      by_key[key] = replacement

  for proj in extracted_projects:
    normalized = normalize_project_entry(proj)
    if not is_valid_project_entry(normalized):
      continue
    key = canonicalize_project_key(normalized.get("name", ""))
    if not key:
      continue
    if key not in by_key:
      by_key[key] = normalized
      order.append(key)
      continue

    current = by_key[key]
    merged_bullets = current.get("bullets", []) + normalized.get("bullets", [])
    unique_bullets = clean_bullet_lines(merged_bullets, max_items=6)

    if not current.get("tech") and normalized.get("tech"):
      current["tech"] = normalized.get("tech")
    current["bullets"] = unique_bullets
    by_key[key] = current

  optimized["projects"] = [by_key[key] for key in order if key in by_key]
  return result


ROLE_HINT_KEYWORDS = {
  "analyst": {
    "analyst", "analysis", "analytics", "sql", "excel", "power bi", "tableau", "dashboard",
    "reporting", "kpi", "cohort", "funnel", "segmentation", "a/b testing", "hypothesis testing",
    "business intelligence", "stakeholder", "ad hoc", "insights",
  },
  "data_scientist": {
    "data scientist", "machine learning", "deep learning", "model", "feature engineering", "xgboost",
    "scikit-learn", "pytorch", "tensorflow", "nlp", "classification", "regression", "forecasting",
    "pyspark", "spark", "databricks", "snowflake", "azure", "gcp", "aws", "big data",
    "price optimization", "demand forecasting", "dimensionality reduction", "clustering", "eda",
    "statistical tests", "model deployment",
  },
  "data_engineer": {
    "data engineer", "etl", "pipeline", "spark", "pyspark", "airflow", "kafka", "dbt", "warehouse",
    "data lake", "bigquery", "redshift", "snowflake", "databricks", "orchestration",
  },
  "software_engineer": {
    "software engineer", "api", "rest", "backend", "microservice", "flask", "fastapi", "django",
    "java", "javascript", "typescript", "system design", "ci/cd", "docker", "kubernetes",
  },
}


def detect_target_role(jd_text):
  text = (jd_text or "").lower()
  if not text:
    return "general"

  role_scores = {}
  for role, hints in ROLE_HINT_KEYWORDS.items():
    score = 0
    for hint in hints:
      if hint in text:
        score += 3 if " " in hint else 1
    role_scores[role] = score

  best_role, best_score = max(role_scores.items(), key=lambda item: item[1])
  return best_role if best_score > 0 else "general"


def score_keyword_matches(text, keywords):
  if not text or not keywords:
    return 0

  lowered = text.lower()
  score = 0
  for keyword in keywords:
    if not keyword:
      continue
    if keyword in lowered:
      score += 3 if " " in keyword else 1
  return score


def score_project_relevance(project, jd_keywords, role_keywords, role="general"):
  if not isinstance(project, dict):
    return 0.0

  name = normalize_space(project.get("name", ""))
  tech = normalize_space(project.get("tech", ""))
  bullets = project.get("bullets", [])
  if not isinstance(bullets, list):
    bullets = []

  corpus = " ".join([name, tech] + [normalize_space(str(b)) for b in bullets])
  jd_score = score_keyword_matches(corpus, jd_keywords)
  role_score = score_keyword_matches(corpus, role_keywords)
  quality_bonus = project_quality_score(project) * 0.35

  role_bonus = 0.0
  lowered = corpus.lower()
  if role == "data_scientist":
    ds_core_terms = [
      "predictive", "forecast", "model deployment", "classification", "regression",
      "clustering", "dimensionality reduction", "spark", "pyspark", "xgboost",
    ]
    ds_hits = sum(1 for term in ds_core_terms if term in lowered)
    role_bonus += ds_hits * 1.2

    # Keep BI helpful but avoid BI-only projects outranking stronger ML projects.
    bi_terms = ["power bi", "dashboard", "kpi design", "reporting"]
    bi_hits = sum(1 for term in bi_terms if term in lowered)
    if ds_hits == 0 and bi_hits > 0:
      role_bonus -= min(2.0, bi_hits * 0.7)

  return (jd_score * 2.0) + (role_score * 1.5) + quality_bonus + role_bonus


def score_experience_relevance(exp, jd_keywords, role_keywords):
  if not isinstance(exp, dict):
    return 0.0

  title = normalize_space(exp.get("title", ""))
  company = normalize_space(exp.get("company", ""))
  bullets = exp.get("bullets", [])
  if not isinstance(bullets, list):
    bullets = []

  corpus = " ".join([title, company] + [normalize_space(str(b)) for b in bullets])
  jd_score = score_keyword_matches(corpus, jd_keywords)
  role_score = score_keyword_matches(corpus, role_keywords)
  bullet_bonus = len([b for b in bullets if normalize_space(str(b))]) * 0.6
  title_bonus = 2.0 if title else 0.0
  return (jd_score * 2.0) + (role_score * 1.2) + bullet_bonus + title_bonus


def apply_role_based_ordering(result, jd_text):
  if not isinstance(result, dict):
    return result

  result = sanitize_resume_entries(result)
  optimized = result.get("optimized_resume")
  if not isinstance(optimized, dict):
    return result

  role = detect_target_role(jd_text)
  role_keywords = list(ROLE_HINT_KEYWORDS.get(role, set()))
  jd_keywords = extract_keywords_from_jd(jd_text, limit=40)

  projects = optimized.get("projects")
  if isinstance(projects, list) and projects:
    scored_projects = []
    for idx, proj in enumerate(projects):
      score = score_project_relevance(proj, jd_keywords, role_keywords, role=role)
      scored_projects.append((score, idx, proj))
    scored_projects.sort(key=lambda item: (-item[0], item[1]))
    optimized["projects"] = [proj for _, _, proj in scored_projects]

  experience = optimized.get("experience")
  if isinstance(experience, list) and experience:
    scored_exp = []
    for idx, exp in enumerate(experience):
      score = score_experience_relevance(exp, jd_keywords, role_keywords)
      scored_exp.append((score, idx, exp))
    scored_exp.sort(key=lambda item: (-item[0], item[1]))
    optimized["experience"] = [exp for _, _, exp in scored_exp]

  return result


TECH_KEYWORDS = {
  # languages
  "python", "sql", "r", "java", "scala", "javascript", "typescript", "bash", "c++",
  # ML/AI
  "machine learning", "deep learning", "nlp", "computer vision", "xgboost",
  "scikit-learn", "tensorflow", "keras", "pytorch", "mllib", "spacy", "transformers",
  "feature engineering", "model deployment", "a/b testing", "statistical modeling",
  "regression", "classification", "clustering", "time series", "forecasting",
  "random forest", "gradient boosting", "neural network", "llm", "generative ai",
  # data eng
  "pyspark", "spark", "hadoop", "kafka", "airflow", "dbt", "etl", "pipeline",
  "data warehouse", "data lake", "bigquery", "redshift", "snowflake", "databricks",
  # visualization
  "power bi", "tableau", "qlik", "looker", "matplotlib", "seaborn", "plotly",
  "streamlit", "dash", "dax", "kpi", "dashboard", "data visualization", "reporting",
  # databases
  "mysql", "postgresql", "mongodb", "redis", "elasticsearch", "sqlite", "nosql",
  # cloud/devops
  "aws", "gcp", "azure", "docker", "kubernetes", "ci/cd", "github actions",
  "ec2", "s3", "lambda", "cloud", "mlops", "prometheus",
  # tools
  "pandas", "numpy", "excel", "alteryx", "jupyter", "git", "flask", "fastapi",
  "rest api", "postman", "jira", "confluence", "notion",
  # analytics domain
  "customer analytics", "campaign analytics", "marketing analytics", "web analytics",
  "cohort analysis", "funnel analysis", "retention", "churn", "segmentation",
  "business intelligence", "performance reporting", "data wrangling", "eda",
  "descriptive statistics", "inferential statistics", "hypothesis testing",
  "data cleaning", "data quality", "kpi design", "ad hoc analysis",
  "stakeholder management", "cross functional", "agile", "data ethics",
}

STOP_WORDS = {
  "with", "that", "from", "this", "have", "your", "will", "were", "been", "into",
  "for", "and", "the", "you", "our", "are", "job", "role", "using", "able", "must",
  "also", "such", "both", "each", "they", "them", "their", "what", "when", "where",
  "which", "while", "about", "after", "before", "between", "through", "during",
  "including", "required", "preferred", "experience", "work", "team", "ability",
  "strong", "good", "excellent", "candidate", "position", "company", "business",
  "help", "build", "make", "need", "use", "new", "per", "well", "etc", "via",
}


def extract_keywords(text, limit=30):
  """Local keyword extractor — fully offline."""
  if not text:
    return []

  tl = text.lower()
  matched = []

  # pass 1: exact phrase matches from domain bank.
  for kw in TECH_KEYWORDS:
    if kw in tl and kw not in matched:
      matched.append(kw)

  # pass 2: high-frequency extra tokens.
  tokens = re.findall(r"[a-z][a-z0-9\+\#\.-]{1,}", tl)
  freq = Counter(t for t in tokens if t not in STOP_WORDS and len(t) > 2)
  ranked_extra = [
    w for w, _ in freq.most_common(60)
    if w not in matched and w not in STOP_WORDS
  ]

  combined = matched + ranked_extra
  return combined[:limit]


def extract_keywords_from_jd(jd_text, resume_text="", limit=30):
  """JD-aware extractor using weighted scoring for bank keywords + frequency."""
  if not jd_text or len(jd_text) < 50:
    return extract_keywords(jd_text or resume_text, limit)

  jd_lower = jd_text.lower()
  scored = {}

  for kw in TECH_KEYWORDS:
    count = jd_lower.count(kw)
    if count > 0:
      scored[kw] = count * 3

  tokens = re.findall(r"[a-z][a-z0-9\+\#\.-]{1,}", jd_lower)
  freq = Counter(t for t in tokens if t not in STOP_WORDS and len(t) > 2)
  for w, c in freq.most_common(80):
    if w not in scored:
      scored[w] = c

  sorted_kws = sorted(scored, key=scored.get, reverse=True)
  return sorted_kws[:limit]


def guess_company_name(jd_text):
  patterns = [
    r"at\s+([A-Z][A-Za-z0-9&.,'\-\s]{2,40})",
    r"company\s*:\s*([A-Z][A-Za-z0-9&.,'\-\s]{2,40})",
  ]
  for pat in patterns:
    m = re.search(pat, jd_text or "")
    if m:
      return normalize_space(m.group(1))
  return "Target Company"


def build_local_fallback_result(resume_text, jd_text):
  resume_lines = [line.strip() for line in (resume_text or "").splitlines() if line.strip()]
  resume_summary_seed = " ".join(resume_lines[:4])[:450]
  jd_keywords = extract_keywords_from_jd(jd_text, resume_text, limit=40)
  resume_keywords = extract_keywords(resume_text, limit=40)
  resume_lower = (resume_text or "").lower()
  matched = [k for k in jd_keywords if k in resume_lower]
  missing = [k for k in jd_keywords if k not in resume_lower][:15]
  match_pct = round((len(matched) / len(jd_keywords)) * 100) if jd_keywords else 0

  header = extract_header_from_resume_text(resume_text)
  certs = extract_certifications_from_resume_text(resume_text)
  company_name = guess_company_name(jd_text)

  skills_primary = sorted({k for k in matched[:18]})
  if not skills_primary:
    skills_primary = resume_keywords[:12]

  role_fit = "Strong" if len(matched) >= 12 else "Moderate" if len(matched) >= 6 else "Low"
  score = min(85, 40 + match_pct // 2)

  return {
    "fallback_mode": True,
    "fallback_reason": "OpenRouter + HuggingFace unavailable — local extraction active",
    "optimized_resume": {
      "name": header.get("name", ""),
      "headline": header.get("headline", ""),
      "summary": (
        "ATS optimization generated in fallback mode due to provider unavailability. "
        + (resume_summary_seed or "Candidate profile extracted from uploaded resume.")
      )[:650],
      "skills": {
        "Programming Languages": [k for k in skills_primary if k.lower() in {"python", "sql", "r", "java", "javascript"}] or skills_primary[:4],
        "Frameworks & Libraries": [k for k in skills_primary if k.lower() in {"tensorflow", "pytorch", "sklearn", "scikit-learn", "flask", "streamlit"}] or skills_primary[4:8],
        "Tools & Technologies": skills_primary[:12],
        "Soft Skills": ["Problem Solving", "Communication", "Collaboration"],
      },
      "experience": [],
      "projects": [],
      "education": [],
      "certifications": certs,
    },
    "ats_score": {
      "total": score,
      "score_reasoning": (
        f"Local mode score based on keyword match only ({match_pct}% match). "
        "Full scoring for bullet quality and structure requires external model APIs."
      ),
    },
    "keyword_analysis": {
      "jd_keywords_extracted": jd_keywords,
      "matched_in_resume": matched,
      "missing_keywords": missing,
      "keyword_match_percentage": match_pct,
    },
    "missing_keywords": missing,
    "improvements": ([
      f"[MISSING KEYWORD] Add '{k}' — found in JD, not in resume"
      for k in missing[:10]
    ] + [
      "Provider fallback mode used due to external API limits; results are conservative.",
      "Quantify impact metrics for top 3 experience bullets.",
    ]),
    "analysis": {
      "role_fit_score": role_fit,
      "matched_keywords": matched,
      "top_strengths": ["Relevant technical keyword overlap", "Resume structured for ATS readability"],
      "weak_points": ["Some JD terms are missing", "Fewer quantified results in bullets"],
      "differentiators": ["Broad tool coverage from resume content"],
    },
    "cover_letter": {
      "company_name": company_name,
      "hiring_manager": "Hiring Manager",
      "subject_line": f"Application for the advertised role — {header.get('name', 'Candidate')}",
      "body": (
        f"Dear Hiring Manager,\n\n"
        f"I am excited to apply for the opportunity at {company_name}. My background aligns with your role requirements and the technical priorities in the job description.\n\n"
        "Across my recent work, I have built practical solutions using the tools listed in my resume, with a strong focus on production readiness, collaboration, and measurable impact.\n\n"
        "I would welcome the chance to discuss how I can contribute to your team and help deliver results from day one.\n\n"
        "Sincerely,\n"
        f"{header.get('name', 'Candidate')}"
      ),
      "word_count": 78,
      "personalization_score": 0,
      "tips": [
        "Add 1-2 company-specific details from the job description to the opening paragraph.",
        "Mention a measurable result from your strongest project or internship experience."
      ],
      "closing": "Sincerely",
      "signature_name": header.get("name", "Candidate"),
    },
    "notice": "External AI providers were unavailable/blocked. Fallback mode generated a usable draft locally.",
  }


def trim_for_model(text, max_chars):
  value = (text or "").strip()
  if len(value) <= max_chars:
    return value, False

  head = value[:max_chars]
  # Keep a natural boundary when possible.
  cut = max(head.rfind("\n"), head.rfind(". "))
  if cut > int(max_chars * 0.7):
    head = head[:cut]
  return head.strip(), True


def extract_first_json_object(text):
  """Extract first balanced JSON object from text while respecting string literals."""
  if not text:
    return None

  start = text.find("{")
  if start == -1:
    return None

  depth = 0
  in_string = False
  escaped = False

  for idx in range(start, len(text)):
    ch = text[idx]
    if escaped:
      escaped = False
      continue

    if ch == "\\":
      escaped = True
      continue

    if ch == '"':
      in_string = not in_string
      continue

    if in_string:
      continue

    if ch == "{":
      depth += 1
    elif ch == "}":
      depth -= 1
      if depth == 0:
        return text[start:idx + 1]

  return None


def parse_ai_json(raw_text):
  """Parse AI response into JSON with common recoveries for wrappers and minor formatting issues."""
  raw = (raw_text or "").strip()
  raw = re.sub(r"^```(?:json)?\s*", "", raw)
  raw = re.sub(r"\s*```$", "", raw)

  try:
    return json.loads(raw)
  except json.JSONDecodeError:
    pass

  candidate = extract_first_json_object(raw)
  if candidate:
    try:
      return json.loads(candidate)
    except json.JSONDecodeError:
      # Remove common trailing comma mistakes.
      cleaned = re.sub(r",\s*([}\]])", r"\1", candidate)
      return json.loads(cleaned)

  # Raise original decoder error shape for caller messaging.
  return json.loads(raw)


def call_anthropic(user_message, system_prompt=SYSTEM_PROMPT):
  response = client.messages.create(
    model="claude-opus-4-5",
    max_tokens=MODEL_MAX_TOKENS,
    system=system_prompt,
    messages=[{"role": "user", "content": user_message}],
  )
  return response.content[0].text.strip()


def call_openrouter(user_message, low_credit_mode=False, api_keys_override=None, system_prompt=SYSTEM_PROMPT):
  url = "https://openrouter.ai/api/v1/chat/completions"
  models_to_try = [OPENROUTER_MODEL] + [m for m in OPENROUTER_MODEL_FALLBACKS if m != OPENROUTER_MODEL]

  def request_with_limit(api_key, model_name, max_tokens):
    payload = {
      "model": model_name,
      "messages": [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": user_message},
      ],
      "temperature": 0.2,
      "max_tokens": max_tokens,
    }

    req = urllib.request.Request(
      url,
      data=json.dumps(payload).encode("utf-8"),
      headers={
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json",
        "HTTP-Referer": "http://localhost:5000",
        "X-Title": "ATS Resume Optimizer",
      },
      method="POST",
    )
    with urllib.request.urlopen(req, timeout=MODEL_TIMEOUT_SECONDS) as res:
      return json.loads(res.read().decode("utf-8"))

  initial_limit = OPENROUTER_LOW_BALANCE_MAX_TOKENS if low_credit_mode else MODEL_MAX_TOKENS
  if api_keys_override is not None:
    api_keys = [key for key in api_keys_override if key]
  else:
    api_keys = [key for key in [OPENROUTER_API_KEY_BACKUP, OPENROUTER_API_KEY] if key]
  if not api_keys:
    raise RuntimeError("No OpenRouter API key configured. Add OPENROUTER_API_KEY to .env")

  body = None
  low_credit_error_seen = False
  endpoint_unavailable_seen = False
  last_error = None

  for model_name in models_to_try:
    for api_key in api_keys:
      try:
        body = request_with_limit(api_key, model_name, initial_limit)
        break
      except urllib.error.HTTPError as http_err:
        details = http_err.read().decode("utf-8", errors="ignore")
        last_error = RuntimeError(f"OpenRouter API error: {http_err.code} {details}")

        if http_err.code == 404 and "No endpoints found" in details:
          endpoint_unavailable_seen = True
          continue

        if http_err.code == 402:
          low_credit_error_seen = True
          try:
            body = request_with_limit(api_key, model_name, OPENROUTER_LOW_BALANCE_MAX_TOKENS)
            break
          except urllib.error.HTTPError as retry_err:
            retry_details = retry_err.read().decode("utf-8", errors="ignore")
            last_error = RuntimeError(f"OpenRouter API error: {retry_err.code} {retry_details}")
            continue

    if body is not None:
      break

  if body is None:
    if low_credit_error_seen:
      raise RuntimeError(
        "OpenRouter credits are low on both primary and backup keys. Recharge credits or shorten JD/resume and retry."
      )
    if endpoint_unavailable_seen:
      raise RuntimeError(
        "OpenRouter model endpoint unavailable. App tried fallback models but none were available."
      )
    if last_error:
      raise last_error
    raise RuntimeError("OpenRouter request failed with both primary and backup keys.")

  choices = body.get("choices", [])
  if not choices:
    raise RuntimeError("OpenRouter returned no choices in response")

  message = choices[0].get("message", {})
  content = message.get("content", "")
  if not content:
    raise RuntimeError("OpenRouter returned an empty message content")

  return content.strip()


def call_huggingface(user_message, system_prompt=SYSTEM_PROMPT):
  api_keys = [key for key in [HUGGINGFACE_API_KEY, HUGGINGFACE_API_KEY_BACKUP] if key]
  if not api_keys:
    raise RuntimeError("No Hugging Face API key configured")

  endpoint = "https://router.huggingface.co/v1/chat/completions"
  models_to_try = [HUGGINGFACE_MODEL] + [m for m in HUGGINGFACE_CHAT_FALLBACK_MODELS if m != HUGGINGFACE_MODEL]

  last_error = None
  for model_name in models_to_try:
    payload = {
      "model": model_name,
      "messages": [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": user_message},
      ],
      "temperature": 0.2,
      "max_tokens": HUGGINGFACE_MAX_NEW_TOKENS,
    }

    for api_key in api_keys:
      req = urllib.request.Request(
        endpoint,
        data=json.dumps(payload).encode("utf-8"),
        headers={
          "Authorization": f"Bearer {api_key}",
          "Content-Type": "application/json",
        },
        method="POST",
      )

      try:
        with urllib.request.urlopen(req, timeout=HUGGINGFACE_TIMEOUT_SECONDS) as res:
          body = json.loads(res.read().decode("utf-8"))
      except urllib.error.HTTPError as http_err:
        details = http_err.read().decode("utf-8", errors="ignore")
        if "<html" in details.lower() or "cloudflare" in details.lower() or "access denied" in details.lower():
          last_error = RuntimeError("Hugging Face API error: 403 provider access blocked by upstream network policy")
        else:
          last_error = RuntimeError(f"Hugging Face API error: {http_err.code} {details}")
        # Try next key/model when model is not supported on chat endpoint.
        if "model_not_supported" in details or "not a chat model" in details:
          continue
        continue

      if isinstance(body, dict):
        choices = body.get("choices", [])
        if choices and isinstance(choices[0], dict):
          message = choices[0].get("message", {})
          content = message.get("content", "")
          if content:
            if isinstance(content, list):
              joined = "".join(
                block.get("text", "") if isinstance(block, dict) else str(block)
                for block in content
              )
              return joined.strip()
            return str(content).strip()

        if body.get("error"):
          last_error = RuntimeError(f"Hugging Face API error: {body.get('error')}")
          continue

      last_error = RuntimeError("Hugging Face returned an unexpected response format")

  if last_error:
    raise last_error
  raise RuntimeError("Hugging Face request failed")


def call_github_models(user_message, system_prompt=SYSTEM_PROMPT):
  if not GITHUB_PAT:
    raise RuntimeError("No GitHub PAT configured")

  endpoint = "https://models.inference.ai.azure.com/chat/completions"
  models_to_try = [GITHUB_MODEL] + [m for m in GITHUB_MODELS_FALLBACKS if m != GITHUB_MODEL]

  last_error = None
  for model_name in models_to_try:
    payload = {
      "model": model_name,
      "messages": [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": user_message},
      ],
      "temperature": 0.2,
      "max_tokens": MODEL_MAX_TOKENS,
    }

    req = urllib.request.Request(
      endpoint,
      data=json.dumps(payload).encode("utf-8"),
      headers={
        "Authorization": f"Bearer {GITHUB_PAT}",
        "Content-Type": "application/json",
      },
      method="POST",
    )

    try:
      with urllib.request.urlopen(req, timeout=MODEL_TIMEOUT_SECONDS) as res:
        body = json.loads(res.read().decode("utf-8"))
    except urllib.error.HTTPError as http_err:
      details = http_err.read().decode("utf-8", errors="ignore")
      last_error = RuntimeError(f"GitHub Models API error: {http_err.code} {details}")
      continue

    choices = body.get("choices", []) if isinstance(body, dict) else []
    if choices and isinstance(choices[0], dict):
      message = choices[0].get("message", {})
      content = message.get("content", "")
      if content:
        if isinstance(content, list):
          joined = "".join(
            block.get("text", "") if isinstance(block, dict) else str(block)
            for block in content
          )
          return joined.strip()
        return str(content).strip()

    last_error = RuntimeError("GitHub Models returned an unexpected response format")

  if last_error:
    raise last_error
  raise RuntimeError("GitHub Models request failed")


def generate_model_response(user_message, low_credit_mode=False, system_prompt=SYSTEM_PROMPT):
  if OPENROUTER_API_KEY_BACKUP:
    try:
      return call_openrouter(user_message, low_credit_mode=low_credit_mode, api_keys_override=[OPENROUTER_API_KEY_BACKUP], system_prompt=system_prompt), "openrouter_backup"
    except Exception:
      pass
  if GITHUB_PAT:
    try:
      return call_github_models(user_message, system_prompt=system_prompt), "github"
    except Exception:
      pass
  if OPENROUTER_API_KEY:
    try:
      return call_openrouter(user_message, low_credit_mode=low_credit_mode, api_keys_override=[OPENROUTER_API_KEY], system_prompt=system_prompt), "openrouter"
    except Exception:
      pass
  if HUGGINGFACE_API_KEY:
    try:
      return call_huggingface(user_message, system_prompt=system_prompt), "huggingface"
    except Exception:
      pass
  if HUGGINGFACE_API_KEY_BACKUP:
    try:
      return call_huggingface(user_message, system_prompt=system_prompt), "huggingface"
    except Exception:
      pass
  if client:
    return call_anthropic(user_message, system_prompt=system_prompt), "anthropic"
  raise RuntimeError(
    "No model API key configured. Add OPENROUTER_API_KEY_BACKUP, GITHUB_PAT, OPENROUTER_API_KEY, HUGGINGFACE_API_KEY, or ANTHROPIC_API_KEY to .env"
  )


def generate_cover_letter_response(resume_text, jd_text, low_credit_mode=False):
  company_name = guess_company_name(jd_text)
  header = extract_header_from_resume_text(resume_text)
  resume_for_model, _ = trim_for_model(resume_text, MODEL_INPUT_MAX_CHARS)
  jd_for_model, _ = trim_for_model(jd_text, MODEL_INPUT_MAX_CHARS)

  user_message = f"""Resume:
{resume_for_model}

Job Description:
{jd_for_model}

Target company: {company_name}
Candidate name: {header.get('name', 'Candidate')}

Write the cover letter now and return only the JSON object defined in the prompt."""

  raw, provider = generate_model_response(user_message, low_credit_mode=low_credit_mode, system_prompt=COVER_LETTER_PROMPT)
  result = parse_ai_json(raw)

  if isinstance(result, dict) and isinstance(result.get("cover_letter"), dict):
    cover_letter = result["cover_letter"]
  else:
    cover_letter = result if isinstance(result, dict) else {}

  normalized_cover = cover_letter
  if provider == "github":
    normalized_cover["notice"] = "GitHub Models was used via GITHUB_PAT."
  elif provider == "huggingface":
    normalized_cover["notice"] = "Hugging Face fallback was used (OpenRouter unavailable/low credits)."
  elif provider == "openrouter_backup":
    normalized_cover["notice"] = "OpenRouter backup key was used first."

  return {
    "cover_letter": normalized_cover,
    "personalization_score": normalized_cover.get("personalization_score", 0),
    "tips": normalized_cover.get("tips", []),
    "provider": provider,
  }


def attach_cover_letter(result, resume_text, jd_text, low_credit_mode=False):
  """Attach a dedicated cover letter to an optimize response."""
  try:
    cover_payload = generate_cover_letter_response(resume_text, jd_text, low_credit_mode=low_credit_mode)
    if isinstance(cover_payload, dict) and isinstance(cover_payload.get("cover_letter"), dict):
      result["cover_letter"] = cover_payload["cover_letter"]
      if cover_payload.get("provider"):
        result["cover_letter_provider"] = cover_payload["provider"]
      return result
  except Exception:
    pass

  if not isinstance(result.get("cover_letter"), dict):
    fallback_cover = build_local_fallback_result(resume_text, jd_text).get("cover_letter", {})
    result["cover_letter"] = fallback_cover
  return result


def extract_resume_text(file_storage):
  filename = (file_storage.filename or "").strip()
  if not filename:
    raise ValueError("Missing file name")

  lower_name = filename.lower()
  file_bytes = file_storage.read()
  if not file_bytes:
    raise ValueError("Uploaded file is empty")
  if len(file_bytes) > MAX_UPLOAD_BYTES:
    max_mb = max(1, MAX_UPLOAD_BYTES // (1024 * 1024))
    raise ValueError(f"Uploaded file is too large. Maximum size is {max_mb} MB")

  if lower_name.endswith(".txt"):
    return file_bytes.decode("utf-8", errors="ignore").strip()

  if lower_name.endswith(".pdf"):
    reader = PdfReader(io.BytesIO(file_bytes))
    extracted = []
    for page in reader.pages:
      extracted.append(page.extract_text() or "")
    return "\n".join(extracted).strip()

  if lower_name.endswith(".docx"):
    doc = Document(io.BytesIO(file_bytes))
    return "\n".join(p.text for p in doc.paragraphs if p.text and p.text.strip()).strip()

  if lower_name.endswith(".doc"):
    raise ValueError(".doc format is not supported. Please upload .docx, .pdf, or .txt")

  raise ValueError("Unsupported file type. Please upload .pdf, .docx, or .txt")


LINK_TOKEN_RE = re.compile(
  r"(https?://\S+|www\.\S+|[\w\.-]+@[\w\.-]+\.\w+|(?:linkedin\.com|github\.com)/\S+)",
  re.IGNORECASE,
)
MARKDOWN_LINK_RE = re.compile(r"\[([^\]]+)\]\((https?://[^)]+)\)")
EMAIL_RE = re.compile(r"[\w\.-]+@[\w\.-]+\.\w+", re.IGNORECASE)
PHONE_RE = re.compile(r"\+?\d[\d\-\s]{7,}\d")
LINKEDIN_RE = re.compile(r"(?:https?://)?(?:www\.)?linkedin\.com/\S+", re.IGNORECASE)
GITHUB_RE = re.compile(r"(?:https?://)?(?:www\.)?github\.com/\S+", re.IGNORECASE)


def normalize_link_target(token):
  cleaned = token.strip()
  if "@" in cleaned and " " not in cleaned and not cleaned.lower().startswith("http"):
    return f"mailto:{cleaned}"
  if cleaned.lower().startswith(("http://", "https://", "mailto:")):
    return cleaned
  if cleaned.lower().startswith("www."):
    return f"https://{cleaned}"
  if cleaned.lower().startswith(("linkedin.com", "github.com")):
    return f"https://{cleaned}"
  return cleaned


def add_hyperlink(paragraph, display_text, url):
  r_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)

  hyperlink = OxmlElement("w:hyperlink")
  hyperlink.set(qn("r:id"), r_id)

  new_run = OxmlElement("w:r")
  run_properties = OxmlElement("w:rPr")

  color = OxmlElement("w:color")
  color.set(qn("w:val"), "0563C1")
  run_properties.append(color)

  underline = OxmlElement("w:u")
  underline.set(qn("w:val"), "single")
  run_properties.append(underline)

  text = OxmlElement("w:t")
  text.text = display_text

  new_run.append(run_properties)
  new_run.append(text)
  hyperlink.append(new_run)
  paragraph._p.append(hyperlink)


def add_text_with_links(paragraph, text, bold=False, italic=False):
  if not text:
    return

  def add_plain_segment(segment):
    parts = LINK_TOKEN_RE.split(segment)
    for part in parts:
      if not part:
        continue

      if LINK_TOKEN_RE.fullmatch(part):
        token = part
        trailing = ""
        while token and token[-1] in ".,);]":
          trailing = token[-1] + trailing
          token = token[:-1]

        if token:
          add_hyperlink(paragraph, token, normalize_link_target(token))
        if trailing:
          run = paragraph.add_run(trailing)
          run.bold = bold
          run.italic = italic
        continue

      run = paragraph.add_run(part)
      run.bold = bold
      run.italic = italic

  cursor = 0
  for md_match in MARKDOWN_LINK_RE.finditer(text):
    add_plain_segment(text[cursor:md_match.start()])
    add_hyperlink(paragraph, md_match.group(1), md_match.group(2))
    cursor = md_match.end()

  add_plain_segment(text[cursor:])


def add_section_heading(doc, title):
  p = doc.add_paragraph()
  run = p.add_run(title.upper())
  run.bold = True
  run.font.size = Pt(12)

  p_format = p.paragraph_format
  p_format.space_before = Pt(10)
  p_format.space_after = Pt(4)

  p_pr = p._p.get_or_add_pPr()
  p_bdr = OxmlElement("w:pBdr")
  bottom = OxmlElement("w:bottom")
  bottom.set(qn("w:val"), "single")
  bottom.set(qn("w:sz"), "6")
  bottom.set(qn("w:space"), "1")
  bottom.set(qn("w:color"), "auto")
  p_bdr.append(bottom)
  p_pr.append(p_bdr)


def normalize_space(value):
  return re.sub(r"\s+", " ", (value or "").strip())


def extract_header_from_resume_text(raw_resume):
  text = raw_resume or ""
  lines = [line.strip() for line in text.splitlines() if line.strip()]

  name = ""
  name_index = -1
  for line in lines[:6]:
    cleaned = normalize_space(line)
    if len(cleaned) < 3 or len(cleaned) > 60:
      continue
    if any(ch.isdigit() for ch in cleaned):
      continue
    if "@" in cleaned or "linkedin" in cleaned.lower() or "github" in cleaned.lower():
      continue
    name = cleaned.upper()
    name_index = lines.index(line)
    break

  headline = ""
  banned_words = (
    "summary",
    "skills",
    "experience",
    "education",
    "projects",
    "objective",
  )
  start_idx = name_index + 1 if name_index >= 0 else 0
  for line in lines[start_idx:start_idx + 5]:
    cleaned = normalize_space(line)
    if len(cleaned) < 6 or len(cleaned) > 90:
      continue
    low = cleaned.lower()
    if "@" in low or "linkedin" in low or "github" in low:
      continue
    if any(word in low for word in banned_words):
      continue
    # Headline usually contains separators like | or role/location words.
    if "|" in cleaned or "scientist" in low or "engineer" in low or "analyst" in low:
      headline = cleaned
      break

  email_match = EMAIL_RE.search(text)
  phone_match = PHONE_RE.search(text)
  linkedin_match = LINKEDIN_RE.search(text)
  github_match = GITHUB_RE.search(text)

  contact_items = []
  if phone_match:
    contact_items.append(normalize_space(phone_match.group(0)))
  if email_match:
    contact_items.append(normalize_space(email_match.group(0)))
  if linkedin_match:
    contact_items.append(normalize_space(linkedin_match.group(0).rstrip(".,;")))
  if github_match:
    contact_items.append(normalize_space(github_match.group(0).rstrip(".,;")))

  return {"name": name, "headline": headline, "contact_items": contact_items}


def set_run_font(run, size=12, bold=False, italic=False):
  run.font.name = "Times New Roman"
  run.font.size = Pt(size)
  run.bold = bold
  run.italic = italic


def style_paragraph(paragraph, spacing_before=0, spacing_after=0, line_spacing=1.0):
  fmt = paragraph.paragraph_format
  fmt.space_before = Pt(spacing_before)
  fmt.space_after = Pt(spacing_after)
  fmt.line_spacing = line_spacing


def add_section_heading_exact(doc, title):
  p = doc.add_paragraph()
  style_paragraph(p, spacing_before=4, spacing_after=2, line_spacing=1.0)
  run = p.add_run(title.upper())
  run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)
  set_run_font(run, size=13, bold=True)

  p_pr = p._p.get_or_add_pPr()
  p_bdr = OxmlElement("w:pBdr")
  bottom = OxmlElement("w:bottom")
  bottom.set(qn("w:val"), "single")
  bottom.set(qn("w:sz"), "6")
  bottom.set(qn("w:space"), "1")
  bottom.set(qn("w:color"), "1A56DB")
  p_bdr.append(bottom)
  p_pr.append(p_bdr)


def add_text_with_links_styled(paragraph, text, size=12, bold=False, italic=False):
  if not text:
    return

  def add_plain_segment(segment):
    parts = LINK_TOKEN_RE.split(segment)
    for part in parts:
      if not part:
        continue

      if LINK_TOKEN_RE.fullmatch(part):
        token = part.rstrip(".,);]")
        trailing = part[len(token):]
        if token:
          add_hyperlink(paragraph, token, normalize_link_target(token))
        if trailing:
          run = paragraph.add_run(trailing)
          set_run_font(run, size=size, bold=bold, italic=italic)
        continue

      run = paragraph.add_run(part)
      set_run_font(run, size=size, bold=bold, italic=italic)

  cursor = 0
  for md_match in MARKDOWN_LINK_RE.finditer(text):
    add_plain_segment(text[cursor:md_match.start()])
    add_hyperlink(paragraph, md_match.group(1), md_match.group(2))
    cursor = md_match.end()

  add_plain_segment(text[cursor:])


def check_domain_alignment(resume_summary, jd_text, low_credit_mode=False):
    job_title = jd_text[:200] if jd_text else "Unknown Job"
    user_message = f"""Job Context:
{job_title}

Resume Summary:
{resume_summary}

Determine the domain alignment penalty."""
    try:
        raw, _ = generate_model_response(user_message, low_credit_mode=low_credit_mode, system_prompt=ALIGNMENT_PROMPT)
        res = parse_ai_json(raw)
        if isinstance(res, dict) and "penalty" in res:
            return min(30, max(0, int(res["penalty"])))
    except Exception:
        pass
    return 0


def compute_hybrid_score(result, jd_text, low_credit_mode=False):
    if not isinstance(result, dict):
        return result
    optimized = result.get("optimized_resume")
    if not isinstance(optimized, dict):
        return result
    
    from utils.ats_scorer import calculate_advanced_ats_score
    
    # Calculate detailed ATS score metrics using the unified scorer
    score_details = calculate_advanced_ats_score(
        resume_text="",
        jd_text=jd_text,
        optimized_resume_dict=optimized,
        low_credit_mode=low_credit_mode
    )
    
    # Save the updated detailed score properties into the response
    result["ats_score"] = {
        "total": score_details["total"],
        "breakdown": score_details["breakdown"],
        "score_reasoning": score_details["score_reasoning"]
    }
    
    # Preserve key properties for direct front-end consumption and compatibility
    result["keyword_analysis"] = score_details["keyword_analysis"]
    result["missing_keywords"] = score_details["keyword_analysis"]["missing_keywords"]
    result["formatting_analysis"] = score_details["formatting_analysis"]
    
    return result


def index():
    return render_template("index.html")


def optimize():
    data = request.get_json(silent=True) or {}
    resume_text = data.get("resume", "").strip()
    jd_text = data.get("jd", "").strip()

    from flask import session
    user_id = session.get('user_id')
    if user_id and resume_text:
        try:
            from models import User
            user = db.session.get(User, user_id)
            if user:
                if user.plan == 'free' and user.api_credits <= 0:
                    return jsonify({
                        "error": "You have exhausted your Free Tier credits. Please upgrade to Pro or Premium Elite to continue!",
                        "limit_reached": True
                    }), 403
                user.resume_text = resume_text
                if user.plan == 'free':
                    user.api_credits = max(0, user.api_credits - 1)
                db.session.commit()
        except Exception:
            db.session.rollback()
    low_credit_mode = bool(data.get("low_credit_mode", False))

    if not resume_text or not jd_text:
        return jsonify({"error": "Both resume and job description are required."}), 400

    resume_for_model, resume_trimmed = trim_for_model(resume_text, MODEL_INPUT_MAX_CHARS)
    jd_for_model, jd_trimmed = trim_for_model(jd_text, MODEL_INPUT_MAX_CHARS)

    user_message = f"""Here is the resume and job description to optimize:

RESUME:
  {resume_for_model}

JOB DESCRIPTION:
  {jd_for_model}

Analyze and optimize the resume for this specific job.
Return ONLY valid JSON, no markdown, no explanation.
Rewrite content for JD fit, but DO NOT remove any original project or certification entries.
Preserve all projects and all certifications from the original resume."""

    try:
      raw, provider = generate_model_response(user_message, low_credit_mode=low_credit_mode)
      result = parse_ai_json(raw)
      result = ensure_certifications_present(result, resume_text)
      result = ensure_projects_present(result, resume_text)
      result = apply_role_based_ordering(result, jd_text)
      result = compute_hybrid_score(result, jd_text, low_credit_mode=low_credit_mode)
      if resume_trimmed or jd_trimmed:
        result["notice"] = "Input was trimmed for speed. If you need full-context optimization, split JD into essentials and rerun."
      if low_credit_mode:
        extra = "Low Credit Mode is enabled; output depth may be shorter to fit token budget."
        result["notice"] = f"{result.get('notice', '')} {extra}".strip()
      if provider == "huggingface":
        extra = "Hugging Face fallback was used (OpenRouter unavailable/low credits)."
        result["notice"] = f"{result.get('notice', '')} {extra}".strip()
      if provider == "github":
        extra = "GitHub Models was used via GITHUB_PAT."
        result["notice"] = f"{result.get('notice', '')} {extra}".strip()
      return jsonify(result)

    except RuntimeError:
        fallback = build_local_fallback_result(resume_text, jd_text)
        fallback = ensure_certifications_present(fallback, resume_text)
        fallback = ensure_projects_present(fallback, resume_text)
        fallback = apply_role_based_ordering(fallback, jd_text)
        fallback = compute_hybrid_score(fallback, jd_text, low_credit_mode=low_credit_mode)
        return jsonify(fallback)

    except json.JSONDecodeError as e:
        # One automatic retry with strict compact-output reminder.
        try:
          retry_message = user_message + "\n\nIMPORTANT: Your previous response was invalid/truncated. Return compact valid JSON only."
          retry_raw, retry_provider = generate_model_response(retry_message, low_credit_mode=low_credit_mode)
          retry_result = parse_ai_json(retry_raw)
          retry_result = ensure_certifications_present(retry_result, resume_text)
          retry_result = ensure_projects_present(retry_result, resume_text)
          retry_result = apply_role_based_ordering(retry_result, jd_text)
          retry_result = compute_hybrid_score(retry_result, jd_text, low_credit_mode=low_credit_mode)
          if resume_trimmed or jd_trimmed:
            retry_result["notice"] = "Input was trimmed for speed. If you need full-context optimization, split JD into essentials and rerun."
          if low_credit_mode:
            extra = "Low Credit Mode is enabled; output depth may be shorter to fit token budget."
            retry_result["notice"] = f"{retry_result.get('notice', '')} {extra}".strip()
          if retry_provider == "huggingface":
            extra = "Hugging Face fallback was used (OpenRouter unavailable/low credits)."
            retry_result["notice"] = f"{retry_result.get('notice', '')} {extra}".strip()
          if retry_provider == "github":
            extra = "GitHub Models was used via GITHUB_PAT."
            retry_result["notice"] = f"{retry_result.get('notice', '')} {extra}".strip()
          return jsonify(retry_result)
        except Exception:
            return jsonify({"error": f"Failed to parse AI response: {str(e)}"}), 500
    except Exception as e:
        return jsonify({"error": str(e)}), 500
BOOST_BULLET_PROMPT = """You are an expert technical resume writer and career coach.
Your task is to take a single resume bullet point, a target role, and optional job description context, and rewrite it into 3 high-impact variations using Google's X-Y-Z formula: "Accomplished [X] as measured by [Y], by doing [Z]".

Rules:
1. You MUST inject realistic, domain-specific, quantifiable metrics (percentages, dollar amounts, time saved, scale, database sizes, or throughput increases).
2. The language must be professional, punchy, and start with strong action verbs (e.g., Spearheaded, Engineered, Orchestrated, Optimized).
3. Align the metrics and technical context specifically with the provided Target Role and Job Description.
4. Keep the results concise and ATS-friendly (maximum 1-2 sentences per bullet).

Respond ONLY with a valid JSON object containing a list of exactly 3 strings in the "variations" field:
{
  "variations": [
    "Accomplished [X] as measured by [Y], by doing [Z]",
    "Accomplished [X] as measured by [Y], by doing [Z]",
    "Accomplished [X] as measured by [Y], by doing [Z]"
  ]
}"""

def build_local_fallback_boost(bullet, role, jd_context):
    role_lower = (role or "general").lower()
    clean_b = bullet.strip(" -•*▸ ")
    
    if any(k in role_lower for k in ["data scientist", "data science", "machine learning", "ml", "ai"]):
        variations = [
            f"Optimized machine learning model accuracy by 14% resulting in $120K annual compute savings, accomplished by engineering custom feature pipelines and hyperparameter grids for {clean_b}.",
            f"Reduced pipeline latency by 35% to process 12M+ records daily, accomplished by refactoring training data ingestion routines and deploying distributed training models based on {clean_b}.",
            f"Boosted predictive precision by 18% as measured by precision/recall metrics, accomplished by implementing ensemble gradient boosting architectures and deploying a monitoring dashboard for {clean_b}."
        ]
    elif any(k in role_lower for k in ["data engineer", "data engineering", "pipeline"]):
        variations = [
            f"Accelerated ETL batch processing throughput by 42%, accomplished by refactoring PySpark joins and building parallel processing queues for {clean_b}.",
            f"Eliminated 99.9% of downstream pipeline failures by implementing automated data quality validations, saving 15 engineering hours weekly while supporting {clean_b}.",
            f"Reduced cloud warehousing query costs by 28% while handling a 3TB database scale, accomplished by optimizing partitioning schemes and dynamic indexing designs for {clean_b}."
        ]
    elif any(k in role_lower for k in ["software", "developer", "backend", "frontend", "engineer"]):
        variations = [
            f"Reduced application server response times by 180ms (40% improvement), accomplished by optimizing SQL query execution plans and implementing redis caching mechanisms for {clean_b}.",
            f"Improved code coverage from 65% to 92% across core microservices, accomplished by authoring 80+ integration test suites and establishing custom CI/CD pipelines to support {clean_b}.",
            f"Scaled microservices infrastructure to handle a 2.5x surge in concurrent requests (up to 15K DAU), accomplished by designing modular RESTful API endpoints and docker container clusters for {clean_b}."
        ]
    else:
        variations = [
            f"Increased reporting accuracy by 25% and reduced ad-hoc data extraction requests by 30 hours monthly, accomplished by designing interactive automated dashboards using {clean_b}.",
            f"Identified $85K in operational inefficiencies, accomplished by executing cohort retention studies and funnel analysis reports based on {clean_b}.",
            f"Enhanced stakeholder decision-making speeds by 3x, accomplished by translating complex data telemetry into weekly executive summaries highlighting {clean_b}."
        ]
    
    return {
        "variations": variations,
        "fallback_mode": True,
        "fallback_reason": "External API keys unavailable or model timed out. Dynamic domain templates generated locally."
    }

def boost_bullet_api():
    data = request.get_json(silent=True) or {}
    bullet = data.get("bullet", "").strip()
    role = data.get("role", "").strip()
    jd_context = data.get("jd_context", "").strip()

    if not bullet:
        return jsonify({"error": "Original bullet text is required."}), 400

    user_message = f"""Original Bullet: "{bullet}"
Target Role: "{role}"
Job Description Context: "{jd_context}"

Generate the 3 optimized variations now using Google's X-Y-Z formula. Return only the JSON object."""

    try:
        raw, provider = generate_model_response(user_message, system_prompt=BOOST_BULLET_PROMPT)
        result = parse_ai_json(raw)
        if not isinstance(result, dict) or "variations" not in result or not isinstance(result["variations"], list):
            raise ValueError("Invalid JSON structure returned by model")
        
        variations = [str(v).strip() for v in result["variations"] if v]
        if len(variations) < 3:
            raise ValueError("Less than 3 variations generated")
        result["variations"] = variations[:3]
        result["provider"] = provider
        return jsonify(result)
    except Exception as e:
        fallback = build_local_fallback_boost(bullet, role, jd_context)
        return jsonify(fallback)


def generate_cover_letter():
    data = request.get_json(silent=True) or {}
    resume_text = data.get("resume", "").strip()
    jd_text = data.get("jd", "").strip()
    low_credit_mode = bool(data.get("low_credit_mode", False))

    if not resume_text or not jd_text:
        return jsonify({"error": "Both resume and job description are required."}), 400

    try:
        payload = generate_cover_letter_response(resume_text, jd_text, low_credit_mode=low_credit_mode)
        return jsonify(payload)
    except Exception as e:
        fallback = build_local_fallback_result(resume_text, jd_text).get("cover_letter", {})
        return jsonify({
            "cover_letter": fallback,
            "personalization_score": fallback.get("personalization_score", 0),
            "tips": fallback.get("tips", []),
            "error": str(e),
        }), 200


def extract_resume():
    uploaded = request.files.get("resume_file")
    if not uploaded:
        return jsonify({"error": "Please choose a file to upload."}), 400

    try:
        from flask import session
        user_id = session.get('user_id')
        if user_id:
            try:
                from models import User
                user = db.session.get(User, user_id)
                if user and user.plan == 'free' and user.api_credits <= 0:
                    return jsonify({
                        "error": "You have exhausted your Free Tier credits. Please upgrade to Pro or Premium Elite to continue!",
                        "limit_reached": True
                    }), 403
            except Exception:
                pass

        text = extract_resume_text(uploaded)
        if not text:
            return jsonify({"error": "Could not extract text from file."}), 400

        if user_id and text:
            try:
                from models import User
                user = db.session.get(User, user_id)
                if user:
                    user.resume_text = text
                    if user.plan == 'free':
                        user.api_credits = max(0, user.api_credits - 1)
                    db.session.commit()
            except Exception:
                db.session.rollback()

        return jsonify({"text": text})
    except ValueError as e:
        return jsonify({"error": str(e)}), 400
    except Exception as e:
        return jsonify({"error": f"Failed to process file: {str(e)}"}), 500


def export_docx():
    data = request.get_json(silent=True) or {}
    resume_data = data.get("resume_data")
    original_resume = data.get("original_resume", "")

    if not resume_data:
        return jsonify({"error": "No resume data provided"}), 400

    doc = Document()
    section = doc.sections[0]
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)

    # Strict serif template to match 10pt letterpaper layout.
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(10)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.line_spacing = 1.0

    try:
        list_style = doc.styles["List Bullet"]
        list_style.font.name = "Times New Roman"
        list_style.font.size = Pt(10)
        list_style.paragraph_format.space_before = Pt(0)
        list_style.paragraph_format.space_after = Pt(0)
        list_style.paragraph_format.line_spacing = 1.0
    except KeyError:
        pass
        
    def add_entry_table(doc, title_text, company_text, right_top, right_bottom=""):
        table = doc.add_table(rows=1, cols=2)
        table.autofit = False
        tblPr = table._tbl.tblPr
        tblW = OxmlElement('w:tblW')
        tblW.set(qn('w:w'), '5000')
        tblW.set(qn('w:type'), 'pct')
        tblPr.append(tblW)
        
        c_left = table.cell(0,0)
        c_right = table.cell(0,1)
        
        p_title = c_left.paragraphs[0]
        style_paragraph(p_title, spacing_before=0, spacing_after=0, line_spacing=1.0)
        if title_text:
            add_text_with_links_styled(p_title, title_text, size=10, bold=True)
            
        if company_text:
            p_comp = c_left.add_paragraph()
            style_paragraph(p_comp, spacing_before=0, spacing_after=0, line_spacing=1.0)
            add_text_with_links_styled(p_comp, company_text, size=10, italic=True)
            
        p_right_top = c_right.paragraphs[0]
        style_paragraph(p_right_top, spacing_before=0, spacing_after=0, line_spacing=1.0)
        p_right_top.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if right_top:
            add_text_with_links_styled(p_right_top, right_top, size=10)
            
        if right_bottom:
            p_right_bot = c_right.add_paragraph()
            style_paragraph(p_right_bot, spacing_before=0, spacing_after=0, line_spacing=1.0)
            p_right_bot.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            add_text_with_links_styled(p_right_bot, right_bottom, size=10, italic=True)
            
        if len(c_left.paragraphs) > 0:
            c_left.paragraphs[-1].paragraph_format.space_after = Pt(2)
        if len(c_right.paragraphs) > 0:
            c_right.paragraphs[-1].paragraph_format.space_after = Pt(2)

    header_from_raw = extract_header_from_resume_text(original_resume)
    contact = resume_data.get("contact", {}) if isinstance(resume_data.get("contact"), dict) else {}

    name_value = normalize_space(
        resume_data.get("name")
        or resume_data.get("full_name")
        or header_from_raw.get("name")
        or "OPTIMIZED RESUME"
    )

    name_para = doc.add_paragraph()
    style_paragraph(name_para, spacing_before=2, spacing_after=2, line_spacing=1.0)
    name_run = name_para.add_run(name_value.upper())
    set_run_font(name_run, size=16, bold=True)
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    headline_value = normalize_space(resume_data.get("headline") or header_from_raw.get("headline") or "")
    if headline_value:
      headline_para = doc.add_paragraph()
      style_paragraph(headline_para, spacing_before=0, spacing_after=1, line_spacing=1.0)
      headline_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
      add_text_with_links_styled(headline_para, headline_value, size=12, bold=True)

    contact_items = []
    for key in ["phone", "email", "linkedin", "github"]:
        value = normalize_space(contact.get(key) or resume_data.get(key) or "")
        if value:
            contact_items.append(value)
    if not contact_items:
        contact_items = header_from_raw.get("contact_items", [])

    if contact_items:
        contact_para = doc.add_paragraph()
        style_paragraph(contact_para, spacing_before=0, spacing_after=3, line_spacing=1.0)
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for index, item in enumerate(contact_items):
            if index > 0:
                sep = contact_para.add_run(" | ")
                set_run_font(sep, size=11)
            add_text_with_links_styled(contact_para, item, size=11)

    add_section_heading_exact(doc, "Professional Summary")
    summary_para = doc.add_paragraph()
    style_paragraph(summary_para, spacing_before=1, spacing_after=4, line_spacing=1.0)
    add_text_with_links_styled(summary_para, resume_data.get("summary", ""), size=10)

    add_section_heading_exact(doc, "Skills")
    skills = resume_data.get("skills", {})
    skill_order = [
        "Skills",
        "Frameworks & Libraries",
        "Programming Languages",
        "Tools & Technologies",
        "Soft Skills",
        "Languages",
    ]
    rendered = set()

    for category in skill_order:
        skill_list = skills.get(category, [])
        if skill_list:
            p = doc.add_paragraph()
            style_paragraph(p, spacing_before=0, spacing_after=2, line_spacing=1.0)
            label = p.add_run(f"{category}: ")
            set_run_font(label, size=10, bold=True)
            add_text_with_links_styled(p, ", ".join(skill_list), size=10)
            rendered.add(category)

    for category, skill_list in skills.items():
        if category in rendered or not skill_list:
            continue
        p = doc.add_paragraph()
        style_paragraph(p, spacing_before=0, spacing_after=2, line_spacing=1.0)
        label = p.add_run(f"{category}: ")
        set_run_font(label, size=10, bold=True)
        add_text_with_links_styled(p, ", ".join(skill_list), size=10)

    experience = resume_data.get("experience", [])
    if experience:
        add_section_heading_exact(doc, "Professional Experience")
        for exp in experience:
            company = normalize_space(exp.get("company", ""))
            title = normalize_space(exp.get("title", ""))
            duration = normalize_space(exp.get("duration", ""))

            location = normalize_space(exp.get("location", ""))
            
            add_entry_table(doc, title_text=title, company_text=company, right_top=duration, right_bottom=location)

            for bullet in exp.get("bullets", []):
                bullet_para = doc.add_paragraph(style="List Bullet")
                style_paragraph(bullet_para, spacing_before=0, spacing_after=0, line_spacing=1.0)
                add_text_with_links_styled(bullet_para, bullet, size=10)

    projects = resume_data.get("projects", [])
    if projects:
        add_section_heading_exact(doc, "Projects")
        for proj in projects:
            project_name = normalize_space(proj.get("name", ""))
            tech = normalize_space(proj.get("tech", ""))
            project_url = normalize_space(proj.get("github") or proj.get("url") or proj.get("link") or "")

            duration = normalize_space(proj.get("duration", "") or proj.get("date", ""))
            
            # Use link if tech is missing, else merge
            company_str = tech
            if project_url and not company_str:
                company_str = project_url
            elif project_url:
                company_str += f" | {project_url}"
                
            add_entry_table(doc, title_text=project_name, company_text=company_str, right_top=duration)

            for bullet in proj.get("bullets", []):
                bullet_para = doc.add_paragraph(style="List Bullet")
                style_paragraph(bullet_para, spacing_before=0, spacing_after=0, line_spacing=1.0)
                add_text_with_links_styled(bullet_para, bullet, size=10)

    education = resume_data.get("education", [])
    if education:
        add_section_heading_exact(doc, "Education")
        for edu in education:
            degree = normalize_space(edu.get("degree", ""))
            institution = normalize_space(edu.get("institution", ""))
            year = normalize_space(edu.get("year", ""))

            location = normalize_space(edu.get("location", ""))
            add_entry_table(doc, title_text=degree, company_text=institution, right_top=year, right_bottom=location)

            details = normalize_space(edu.get("details", ""))
            if details:
                details_para = doc.add_paragraph()
                style_paragraph(details_para, spacing_before=0, spacing_after=0, line_spacing=1.0)
                add_text_with_links_styled(details_para, details, size=10)

        certifications = resume_data.get("certifications", [])
        if False and certifications:
          add_section_heading_exact(doc, "Certifications")
          for cert in certifications:
            if isinstance(cert, str):
              p_cert = doc.add_paragraph()
              style_paragraph(p_cert, spacing_before=1, spacing_after=0, line_spacing=1.0)
              add_text_with_links_styled(p_cert, cert, size=12)
              continue

            name = normalize_space(cert.get("name", ""))
            issuer = normalize_space(cert.get("issuer", ""))
            year = normalize_space(cert.get("year", ""))
            line = name
            if issuer:
              line += f" — {issuer}"
            if year:
              line += f" ({year})"
            if line:
              p_cert = doc.add_paragraph()
              style_paragraph(p_cert, spacing_before=1, spacing_after=0, line_spacing=1.0)
              add_text_with_links_styled(p_cert, line, size=12)

    certifications = resume_data.get("certifications", [])
    if certifications:
        add_section_heading_exact(doc, "Certifications")
        for cert in certifications:
            if isinstance(cert, str):
                p_cert = doc.add_paragraph()
                style_paragraph(p_cert, spacing_before=1, spacing_after=0, line_spacing=1.0)
                add_text_with_links_styled(p_cert, cert, size=10)
                continue

            name = normalize_space(cert.get("name", ""))
            issuer = normalize_space(cert.get("issuer", ""))
            year = normalize_space(cert.get("year", ""))
            line = name
            if issuer:
                line += f" - {issuer}"
            if year:
                line += f" ({year})"
            if line:
                p_cert = doc.add_paragraph()
                style_paragraph(p_cert, spacing_before=1, spacing_after=0, line_spacing=1.0)
                add_text_with_links_styled(p_cert, line, size=10)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return send_file(
        buffer,
        as_attachment=True,
        download_name="optimized_resume.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


def export_cover_letter_docx():
    data = request.get_json(silent=True) or {}
    cover_letter = data.get("cover_letter")
    resume_data = data.get("resume_data") or {}
    original_resume = data.get("original_resume", "")

    if not isinstance(cover_letter, dict):
        return jsonify({"error": "No cover letter data provided"}), 400

    header_from_raw = extract_header_from_resume_text(original_resume)
    candidate_name = normalize_space(
        cover_letter.get("signature_name")
        or resume_data.get("name")
        or resume_data.get("full_name")
        or header_from_raw.get("name")
        or "Candidate"
    )
    company_name = normalize_space(cover_letter.get("company_name", ""))
    hiring_manager = normalize_space(cover_letter.get("hiring_manager") or "Hiring Manager")
    subject = normalize_space(cover_letter.get("subject", ""))
    closing = normalize_space(cover_letter.get("closing") or "Sincerely")

    body_paragraphs = cover_letter.get("body_paragraphs")
    if not body_paragraphs:
        body_paragraphs = cover_letter.get("body", [])
        
    if isinstance(body_paragraphs, str):
        body_paragraphs = [p.strip() for p in re.split(r"\n\s*\n", body_paragraphs) if p.strip()]
    if not isinstance(body_paragraphs, list):
        body_paragraphs = []

    doc = Document()
    section = doc.sections[0]
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    p_date = doc.add_paragraph(date.today().strftime("%B %d, %Y"))
    style_paragraph(p_date, spacing_after=6)

    greeting_target = hiring_manager
    if company_name:
        p_company = doc.add_paragraph(company_name)
        style_paragraph(p_company, spacing_after=1)
    p_greet = doc.add_paragraph(f"Dear {greeting_target},")
    style_paragraph(p_greet, spacing_after=6)

    if subject:
        p_subject = doc.add_paragraph()
        run_subject = p_subject.add_run(f"Subject: {subject}")
        set_run_font(run_subject, size=12, bold=True)
        style_paragraph(p_subject, spacing_after=6)

    for para in body_paragraphs:
        p_body = doc.add_paragraph()
        style_paragraph(p_body, spacing_after=6, line_spacing=1.15)
        add_text_with_links_styled(p_body, para, size=12)

    p_close = doc.add_paragraph(closing + ",")
    style_paragraph(p_close, spacing_before=8, spacing_after=0)
    p_name = doc.add_paragraph(candidate_name)
    style_paragraph(p_name, spacing_after=0)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return send_file(
        buffer,
        as_attachment=True,
        download_name="cover_letter.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

def login():
    return render_template("login.html")

def create_app():
  app = Flask(__name__, static_folder="static", template_folder="templates")

  # Core config
  db_url = os.environ.get('DATABASE_URL')
  if db_url:
    if db_url.startswith("postgres://"):
      db_url = db_url.replace("postgres://", "postgresql://", 1)
    app.config['SQLALCHEMY_DATABASE_URI'] = db_url
    app.config['SQLALCHEMY_ENGINE_OPTIONS'] = {
      'pool_pre_ping': True,
      'pool_recycle': 300
    }
  else:
    from tracker.config import get_db_path
    db_path = get_db_path(app.root_path)
    abs_db_path = os.path.abspath(db_path)
    os.makedirs(os.path.dirname(abs_db_path), exist_ok=True)
    app.config['SQLALCHEMY_DATABASE_URI'] = f'sqlite:///{abs_db_path}'

  app.secret_key = os.environ.get("SECRET_KEY", "dev_fallback_secret_key_123")

  try:
    max_bytes = int(os.getenv("MAX_UPLOAD_BYTES", str(8 * 1024 * 1024)))
  except ValueError:
    max_bytes = 8 * 1024 * 1024
  app.config["MAX_CONTENT_LENGTH"] = max_bytes

  # Initialize extensions
  db.init_app(app)
  bcrypt.init_app(app)
  login_manager.init_app(app)

  @login_manager.user_loader
  def _load_user(user_id):
    try:
      return db.session.get(User, int(user_id))
    except Exception:
      return None

  # Register blueprints
  try:
    from auth.routes import auth_blueprint
    app.register_blueprint(auth_blueprint)
  except Exception:
    pass

  try:
    from tracker.routes import tracker_blueprint
    app.register_blueprint(tracker_blueprint)
  except Exception:
    pass

  try:
    from optimizer.routes import optimizer_blueprint
    app.register_blueprint(optimizer_blueprint)
  except Exception:
    pass

  try:
    from latex_resume import latex_blueprint
    app.register_blueprint(latex_blueprint)
  except Exception:
    pass

  try:
    from billing.routes import billing_blueprint
    app.register_blueprint(billing_blueprint)
  except Exception as e:
    print("Failed to register billing blueprint:", e)

  try:
    from interview.routes import interview_blueprint
    app.register_blueprint(interview_blueprint)
  except Exception as e:
    print("Failed to register interview blueprint:", e)

  with app.app_context():
    db.create_all()
    # Dynamic Auto-Migrations and Seeding
    try:
      from sqlalchemy import text
      # 1. Create boards table if it doesn't exist (failsafe SQL)
      try:
        db.session.execute(text("""
          CREATE TABLE IF NOT EXISTS boards (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id INTEGER NOT NULL,
            name VARCHAR(150) NOT NULL,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY(user_id) REFERENCES users(id) ON DELETE CASCADE
          )
        """))
        db.session.commit()
      except Exception:
        db.session.rollback()  # Rollback aborted SQLite transaction before attempting PostgreSQL fallback
        try:
          db.session.execute(text("""
            CREATE TABLE IF NOT EXISTS boards (
              id SERIAL PRIMARY KEY,
              user_id INTEGER NOT NULL REFERENCES users(id) ON DELETE CASCADE,
              name VARCHAR(150) NOT NULL,
              created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
          """))
          db.session.commit()
        except Exception as e:
          db.session.rollback()
          print("Failed to create boards table:", e)

      # Create interview_sessions table if it doesn't exist (failsafe SQL)
      try:
        db.session.execute(text("""
          CREATE TABLE IF NOT EXISTS interview_sessions (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id INTEGER NOT NULL,
            job_title VARCHAR(150) NOT NULL,
            interview_type VARCHAR(50) DEFAULT 'Mixed',
            difficulty VARCHAR(50) DEFAULT 'Mid',
            score INTEGER,
            feedback TEXT,
            history TEXT,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY(user_id) REFERENCES users(id) ON DELETE CASCADE
          )
        """))
        db.session.commit()
      except Exception:
        db.session.rollback()
        try:
          db.session.execute(text("""
            CREATE TABLE IF NOT EXISTS interview_sessions (
              id SERIAL PRIMARY KEY,
              user_id INTEGER NOT NULL REFERENCES users(id) ON DELETE CASCADE,
              job_title VARCHAR(150) NOT NULL,
              interview_type VARCHAR(50) DEFAULT 'Mixed',
              difficulty VARCHAR(50) DEFAULT 'Mid',
              score INTEGER,
              feedback TEXT,
              history TEXT,
              created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
          """))
          db.session.commit()
        except Exception as e:
          db.session.rollback()
          print("Failed to create interview_sessions table:", e)

      # 2. Add columns to applications table dynamically
      for col_def in [
        "board_id INTEGER DEFAULT 1",
        "archived BOOLEAN DEFAULT FALSE"
      ]:
        try:
          db.session.execute(text(f"ALTER TABLE applications ADD COLUMN {col_def}"))
          db.session.commit()
        except Exception:
          db.session.rollback()

      # Add resume_text column to users table dynamically
      try:
        db.session.execute(text("ALTER TABLE users ADD COLUMN resume_text TEXT"))
        db.session.commit()
      except Exception:
        db.session.rollback()

      # Add billing columns to users table dynamically
      for col_def in [
        "plan VARCHAR(50) DEFAULT 'free'",
        "stripe_customer_id VARCHAR(255)",
        "subscription_active BOOLEAN DEFAULT FALSE",
        "api_credits INTEGER DEFAULT 2"
      ]:
        try:
          db.session.execute(text(f"ALTER TABLE users ADD COLUMN {col_def}"))
          db.session.commit()
        except Exception:
          db.session.rollback()

      # 3. Seed default board for each user if none exists
      try:
        from models import User
        users = User.query.all()
        for u in users:
          first_board = db.session.execute(
            text("SELECT id FROM boards WHERE user_id = :uid LIMIT 1"),
            {"uid": u.id}
          ).fetchone()
          if not first_board:
            db.session.execute(
              text("INSERT INTO boards (user_id, name) VALUES (:uid, :name)"),
              {"uid": u.id, "name": "Standard Hunt"}
            )
            db.session.commit()
      except Exception as e:
        db.session.rollback()
        print("Failed to seed default board:", e)
    except Exception as exc:
      print("Migration pipeline error:", exc)

  return app


app = create_app()


if __name__ == "__main__":
  debug_raw = os.getenv("FLASK_DEBUG", "0").strip().lower()
  debug_enabled = debug_raw in {"1", "true", "yes", "on"}
  try:
    port = int(os.getenv("PORT", "5000"))
  except ValueError:
    port = 5000
  app.run(host="0.0.0.0", debug=debug_enabled, port=port)


