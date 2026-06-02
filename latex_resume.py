import io
import os
import re
import shutil
import subprocess
import tempfile
import textwrap

from flask import Blueprint, jsonify, request, send_file
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

latex_blueprint = Blueprint("latex_resume", __name__)
TEMPLATE_PATH = os.path.join(os.path.dirname(__file__), "resume_template.tex")


# ── engine detection ──────────────────────────────────────────────────────────

def get_latex_engines():
    return [e for e in ["pdflatex", "xelatex", "lualatex"] if shutil.which(e)]


def is_engine_usable(engine):
    if not engine:
        return False
    try:
        r = subprocess.run([engine, "--version"], capture_output=True, text=True, timeout=10, check=False)
        return r.returncode == 0
    except Exception:
        return False


def find_usable_engines():
    return [e for e in get_latex_engines() if is_engine_usable(e)]


# ── error parsing ─────────────────────────────────────────────────────────────

def extract_missing_packages(log):
    if not log:
        return []
    names = set()
    for m in re.findall(r"([A-Za-z0-9_-]+)\.sty", log):
        names.add(m.strip())
    for m in re.findall(r"\\RequirePackage(?:\[[^\]]*\])?\{([^}]+)\}", log):
        names.add(m.strip().split(",")[0].strip())
    return sorted(names)


def last_lines(text, n=15):
    return "\n".join([line for line in (text or "").splitlines() if line.strip()][-n:])


def latex_escape(value):
    text = str(value or "")
    replacements = {
        "\\": r"\textbackslash{}",
        "&": r"\&",
        "%": r"\%",
        "$": r"\$",
        "#": r"\#",
        "_": r"\_",
        "{": r"\{",
        "}": r"\}",
        "~": r"\textasciitilde{}",
        "^": r"\textasciicircum{}",
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    return text


def load_template_text(template_id="classic"):
    filename = "resume_template.tex"
    if template_id == "modern":
        filename = "resume_template_modern.tex"
    elif template_id == "executive":
        filename = "resume_template_executive.tex"
    
    path = os.path.join(os.path.dirname(__file__), filename)
    if not os.path.exists(path):
        path = TEMPLATE_PATH
        
    with open(path, "r", encoding="utf-8") as f:
        return f.read()


def bullet_items(lines):
    items = []
    for line in lines or []:
        clean = str(line or "").strip()
        if clean:
            items.append(f"  \\item {latex_escape(clean)}")
    return "\n".join(items)


def trim_text(value, max_len):
    text = str(value or "").strip()
    if len(text) <= max_len:
        return text
    return text[:max_len].rstrip() + "..."


def trim_lines(lines, max_lines, max_len):
    out = []
    for line in (lines or [])[:max_lines]:
        clean = str(line or "").strip()
        if clean:
            out.append(trim_text(clean, max_len))
    return out


def parse_field_lines(value):
    if isinstance(value, list):
        return [str(item).strip() for item in value if str(item).strip()]
    text = str(value or "")
    return [line.strip() for line in text.splitlines() if line.strip()]


def parse_block_items(value, min_fields, defaults):
    """Parse textarea lines into structured rows.

    Rules:
    - Lines with enough `|` columns start a new row.
    - Plain lines are appended as bullets/details to the current row.
    - If no row exists yet, create one using the first plain line as title/name.
    """
    items = []
    for raw_line in parse_field_lines(value):
        line = str(raw_line or "").strip()
        if not line:
            continue

        parts = [part.strip() for part in line.split("|") if part.strip()]
        if len(parts) >= min_fields:
            items.append(parts)
            continue

        detail = re.sub(r"^[\-•\*\s]+", "", line).strip()
        if not detail:
            continue

        if not items:
            seed = [detail] + list(defaults)
            items.append(seed[:min_fields])
        else:
            items[-1].append(detail)
    return items


MONTH_PATTERN = r"(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec|January|February|March|April|June|July|August|September|October|November|December)"


def clean_detail_line(line):
    return re.sub(r"^[\-•\*\s]+", "", str(line or "")).strip()


def split_title_duration(line):
    text = str(line or "").strip()
    month_range_re = re.compile(
        rf"(?P<dur>{MONTH_PATTERN}\s+\d{{4}}\s*[-–—]\s*(?:Present|{MONTH_PATTERN}\s+\d{{4}}))",
        re.IGNORECASE,
    )
    year_range_re = re.compile(r"(?P<dur>\b(?:19|20)\d{2}\s*[-–—]\s*(?:Present|(?:19|20)\d{2})\b)", re.IGNORECASE)

    m = month_range_re.search(text) or year_range_re.search(text)
    if not m:
        return text, "Date"

    duration = m.group("dur").strip()
    title = (text[:m.start()] + " " + text[m.end():]).strip(" |-:—")
    return (title or text), duration


def is_probable_date_line(line):
    text = str(line or "").strip()
    if not text:
        return False
    month_range_re = re.compile(
        rf"^{MONTH_PATTERN}\s+\d{{4}}\s*[-–—]\s*(?:Present|{MONTH_PATTERN}\s+\d{{4}})$",
        re.IGNORECASE,
    )
    year_range_re = re.compile(r"^(?:19|20)\d{2}\s*[-–—]\s*(?:Present|(?:19|20)\d{2})$", re.IGNORECASE)
    single_year_re = re.compile(r"^(?:19|20)\d{2}$")
    return bool(month_range_re.match(text) or year_range_re.match(text) or single_year_re.match(text))


def is_probable_entry_header(line):
    text = str(line or "").strip()
    if not text:
        return False
    if "|" in text:
        return True
    if re.search(r"\b(?:19|20)\d{2}\b", text):
        return True
    if " - " in text or " — " in text:
        return True
    return False


def parse_experience_entries(value):
    entries = []
    for raw in parse_field_lines(value):
        line = clean_detail_line(raw)
        if not line:
            continue

        if is_probable_date_line(line):
            if entries:
                current = entries[-1]
                if current.get("duration") in {"", "Date"}:
                    current["duration"] = line
                elif current.get("duration") == line:
                    pass
                else:
                    current.setdefault("bullets", []).append(line)
            continue

        parts = [part.strip() for part in line.split("|") if part.strip()]
        if len(parts) >= 4:
            entries.append({
                "title": parts[0],
                "company": parts[1],
                "duration": parts[2],
                "location": parts[3],
                "bullets": parts[4:],
            })
            continue

        if is_probable_entry_header(line):
            title, duration = split_title_duration(line)
            entries.append({
                "title": title,
                "company": "Company",
                "duration": duration,
                "location": "Location",
                "bullets": [],
            })
            continue

        if not entries:
            entries.append({
                "title": line,
                "company": "Company",
                "duration": "Date",
                "location": "Location",
                "bullets": [],
            })
            continue

        current = entries[-1]
        looks_like_company = False
        if current["company"] == "Company":
            words = line.split()
            has_company_markers = bool(re.search(r"\b(pvt\.?|ltd\.?|llp|inc\.?|corp\.?|company|solutions|technologies|systems|labs|group|studio|remote)\b", line, re.IGNORECASE))
            sentence_like = bool(re.search(r"[.;]$", line)) or line[:1].islower() or len(words) > 12
            title_like = bool(is_probable_date_line(line) or is_probable_entry_header(line))
            looks_like_company = (has_company_markers or (not sentence_like and not title_like and len(words) <= 14))

        if looks_like_company:
            loc_match = re.search(r"\(([^)]+)\)", line)
            if loc_match:
                current["location"] = loc_match.group(1).strip()
            company = re.sub(r"\([^)]*\)", "", line).strip(" -")
            current["company"] = company or current["company"]
        else:
            current["bullets"].append(line)

    return entries


def parse_project_entries(value):
    entries = []
    for raw in parse_field_lines(value):
        line = clean_detail_line(raw)
        if not line:
            continue

        if is_probable_date_line(line):
            if entries:
                current = entries[-1]
                if current.get("date") in {"", "Date"}:
                    current["date"] = line
                elif current.get("date") == line:
                    pass
                else:
                    current.setdefault("bullets", []).append(line)
            continue

        parts = [part.strip() for part in line.split("|") if part.strip()]
        if len(parts) >= 3:
            entries.append({
                "name": parts[0],
                "tech": parts[1],
                "date": parts[2],
                "bullets": parts[3:],
            })
            continue

        if is_probable_entry_header(line):
            name, date = split_title_duration(line)
            entries.append({
                "name": name,
                "tech": "Tech Stack",
                "date": date,
                "bullets": [],
            })
            continue

        if not entries:
            entries.append({
                "name": line,
                "tech": "Tech Stack",
                "date": "Date",
                "bullets": [],
            })
            continue

        current = entries[-1]
        looks_like_tech = (
            current["tech"] == "Tech Stack"
            and len(line) <= 140
            and ("|" in line or "," in line)
            and not line.endswith(".")
        )
        if looks_like_tech:
            current["tech"] = line.replace("|", ", ")
        else:
            current["bullets"].append(line)

    return entries


def build_contact_lines(data):
    parts = []
    location = latex_escape(data.get("location") or "City, State")
    phone = latex_escape(data.get("phone") or "+91 XXXXXXXXXX")
    email = latex_escape(data.get("email") or "your@email.com")
    linkedin = latex_escape(data.get("linkedin") or "linkedin.com/in/yourprofile")
    github = latex_escape(data.get("github") or "github.com/yourgithub")

    parts.append(rf"\faMapMarker\ {location}")
    parts.append(rf"\faPhone\ {phone}")
    parts.append(rf"\href{{mailto:{email}}}{{\faEnvelope\ {email}}}")
    parts.append(rf"\href{{https://{linkedin}}}{{\faLinkedin\ {linkedin}}}")
    parts.append(rf"\href{{https://{github}}}{{\faGithub\ {github}}}")
    return "\n    \\quad\\textbar\\quad\n    ".join(parts)


def render_template_document(data, template_id="classic"):
    template = load_template_text(template_id)
    full_name = latex_escape(data.get("full_name") or data.get("name") or "YOUR FULL NAME")
    header = build_contact_lines(data)

    summary = latex_escape(trim_text(data.get("summary") or "Add a concise recruiter-friendly summary here.", 700))

    skills = data.get("skills") or {}
    skill_lines = []
    if isinstance(skills, list) and skills:
        for line in skills:
            text = str(line or "").strip()
            if not text:
                continue
            if ":" in text:
                category, values = text.split(":", 1)
                category = category.strip()
                values = values.strip()
                if category and values:
                    skill_lines.append(rf"  \item \textbf{{{latex_escape(category)}:}} {latex_escape(values)}")
            else:
                skill_lines.append(rf"  \item {latex_escape(text)}")
    elif isinstance(skills, dict) and skills:
        for category, values in skills.items():
            items = ", ".join(latex_escape(item) for item in parse_field_lines(values))
            if items:
                skill_lines.append(rf"  \item \textbf{{{latex_escape(category)}:}} {items}")
    else:
        skill_lines = [r"  \item \textbf{Languages:} Python, SQL"]
        skill_lines.append(r"  \item \textbf{ML / AI:} Scikit-learn, XGBoost")

    skill_lines = skill_lines[:10]

    experience_blocks = []
    for entry in parse_experience_entries(data.get("experience"))[:6]:
        title = latex_escape(entry.get("title") or "Role")
        company = latex_escape(entry.get("company") or "Company")
        duration = latex_escape(entry.get("duration") or "Date")
        location = latex_escape(entry.get("location") or "Location")
        bullets = trim_lines(entry.get("bullets") or [], max_lines=8, max_len=210)
        block = f"\\entry{{{title}}}{{{company}}}{{{duration}}}{{{location}}}"
        bullet_text = bullet_items(bullets)
        if bullet_text:
            block += f"\n\\begin{{itemize}}\n{bullet_text}\n\\end{{itemize}}"
        experience_blocks.append(block)
    if not experience_blocks:
        experience_blocks.append(
            r"\entry{Role}{Company}{Date}{Location}" + "\n" + r"\begin{itemize}" + "\n" + r"  \item Add experience bullets" + "\n" + r"\end{itemize}"
        )

    project_blocks = []
    for entry in parse_project_entries(data.get("projects"))[:6]:
        name = latex_escape(entry.get("name") or "Project")
        tech = latex_escape(entry.get("tech") or "Tech Stack")
        date = latex_escape(entry.get("date") or "Date")
        bullets = trim_lines(entry.get("bullets") or [], max_lines=7, max_len=205)
        block = f"\\project{{{name}}}{{{tech}}}{{{date}}}"
        bullet_text = bullet_items(bullets)
        if bullet_text:
            block += f"\n\\begin{{itemize}}\n{bullet_text}\n\\end{{itemize}}"
        project_blocks.append(block)
    if not project_blocks:
        project_blocks.append(
            r"\project{Project}{Tech Stack}{Date}" + "\n" + r"\begin{itemize}" + "\n" + r"  \item Add project bullets" + "\n" + r"\end{itemize}"
        )

    education_entries = parse_block_items(data.get("education"), min_fields=3, defaults=["Institution", "Year"])
    education_blocks = []
    for entry in education_entries:
        degree = latex_escape(entry[0] if len(entry) > 0 else "Degree")
        institution = latex_escape(entry[1] if len(entry) > 1 else "Institution")
        year = latex_escape(entry[2] if len(entry) > 2 else "Year")
        location = latex_escape(entry[3] if len(entry) > 3 else "Location")
        education_blocks.append(
            f"\\entry{{{degree}}}{{{institution}}}{{{year}}}{{{location}}}"
        )
    if not education_blocks:
        education_blocks.append(r"\entry{Degree}{Institution}{Year}{Location}")

    certifications_lines = trim_lines(parse_field_lines(data.get("certifications")), max_lines=12, max_len=190)
    certifications_block = "\n".join(rf"  \item {latex_escape(item)}" for item in certifications_lines) if certifications_lines else r"  \item Add certifications here"

    rendered = template
    # Replace the name placeholder directly
    rendered = rendered.replace(r"{\LARGE\bfseries YOUR FULL NAME}", "{\\LARGE\\bfseries " + full_name + "}")
    # Replace the contact/small block
    rendered = re.sub(
        r"\{\\small\n(    \\faMapMarker.*?)\n  \}",
        lambda _: "{\\small\n    " + header + "\n  }",
        rendered,
        flags=re.DOTALL,
        count=1,
    )
    rendered = re.sub(
        r"\\section\{Professional Summary\}\n(?:.|\n)*?\n% ══ SKILLS ══",
        lambda _: "\\section{Professional Summary}\n" + summary + "\n\n% ══ SKILLS ══",
        rendered,
        count=1,
    )
    rendered = re.sub(
        r"\\section\{Technical Skills\}\n(?:.|\n)*?\n% ══ EXPERIENCE ══",
        lambda _: "\\section{Technical Skills}\n\\begin{itemize}[leftmargin=0pt, label={}]\n" + "\n".join(skill_lines) + "\n\\end{itemize}\n\n% ══ EXPERIENCE ══",
        rendered,
        count=1,
    )
    rendered = re.sub(
        r"\\section\{Experience\}\n(?:.|\n)*?\n% ══ PROJECTS ══",
        lambda _: "\\section{Experience}\n" + "\n\n".join(experience_blocks) + "\n\n% ══ PROJECTS ══",
        rendered,
        count=1,
    )
    rendered = re.sub(
        r"\\section\{Projects\}\n(?:.|\n)*?\n% ══ EDUCATION ══",
        lambda _: "\\section{Projects}\n" + "\n\n".join(project_blocks) + "\n\n% ══ EDUCATION ══",
        rendered,
        count=1,
    )
    rendered = re.sub(
        r"\\section\{Education\}\n(?:.|\n)*?\n% ══ CERTIFICATIONS \\& COURSES ══",
        lambda _: "\\section{Education}\n" + "\n\n".join(education_blocks) + "\n\n% ══ CERTIFICATIONS \\& COURSES ══",
        rendered,
        count=1,
    )
    rendered = re.sub(
        r"\\section\{Certifications \\& Courses\}\n(?:.|\n)*?\n\\end\{document\}",
        lambda _: "\\section{Certifications \\& Courses}\n\\begin{itemize}\n" + certifications_block + "\n\\end{itemize}\n\n\\end{document}",
        rendered,
        count=1,
    )
    return rendered



@latex_blueprint.route("/render-latex-source", methods=["POST"])
def render_latex_source():
    data = request.get_json(silent=True) or {}
    latex_data = data.get("latex_data") or {}
    template_id = data.get("template_id") or "classic"

    if not isinstance(latex_data, dict) or not latex_data:
        return jsonify({"error": "latex_data is required"}), 400

    try:
        source = render_template_document(latex_data, template_id=template_id)
        source = sanitize_tex(source)
        return jsonify({"source": source})
    except Exception as exc:
        return jsonify({"error": f"Template render failed: {exc}"}), 400


# ── compilation ───────────────────────────────────────────────────────────────

def compile_once(engine, source):
    """Run engine twice (for cross-refs). Returns (pdf_bytes | None, log_str)."""
    with tempfile.TemporaryDirectory(prefix="lx_") as d:
        tex = os.path.join(d, "resume.tex")
        pdf = os.path.join(d, "resume.pdf")
        with open(tex, "w", encoding="utf-8") as f:
            f.write(source)

        cmd = [engine, "-no-shell-escape", "-interaction=nonstopmode", "-halt-on-error", "-file-line-error", "resume.tex"]

        # two passes so \label / \ref resolve
        for _ in range(2):
            r = subprocess.run(cmd, cwd=d, capture_output=True, text=True, timeout=90, check=False)

        if r.returncode != 0 or not os.path.exists(pdf):
            return None, (r.stdout or "") + "\n" + (r.stderr or "")

        with open(pdf, "rb") as f:
            return f.read(), ""


def compile_latex(source):
    engines = find_usable_engines()
    if not engines:
        raise FileNotFoundError(
            "No LaTeX engine found. Install TeX Live (pdflatex) or MiKTeX and ensure it is on PATH."
        )

    errors = []
    for eng in engines:
        pdf, log = compile_once(eng, source)
        if pdf:
            return pdf
        errors.append(f"[{eng}] {last_lines(log)}")

    raise RuntimeError("LaTeX compilation failed:\n" + "\n---\n".join(errors))


# ── plain-text fallback ───────────────────────────────────────────────────────

def latex_to_plain(source):
    t = str(source or "")
    t = re.sub(r"%[^\n]*", "", t)
    t = re.sub(r"\\\\", "\n", t)
    t = re.sub(r"\\documentclass(?:\[[^\]]*\])?\{[^}]+\}", "", t)
    t = re.sub(r"\\usepackage(?:\[[^\]]*\])?\{[^}]+\}", "", t)
    t = re.sub(r"\\begin\{[^}]+\}", "", t)
    t = re.sub(r"\\end\{[^}]+\}", "", t)
    t = re.sub(r"\\section\*?\{([^}]*)\}", r"\n\1\n" + "=" * 40, t)
    t = re.sub(r"\\subsection\*?\{([^}]*)\}", r"\n\1\n", t)
    t = re.sub(r"\\textbf\{([^}]*)\}", r"\1", t)
    t = re.sub(r"\\textit\{([^}]*)\}", r"\1", t)
    t = re.sub(r"\\href\{[^}]*\}\{([^}]*)\}", r"\1", t)
    t = re.sub(r"\\hfill", "    ", t)
    t = re.sub(r"\$\\vert\$", " | ", t)
    t = re.sub(r"\\[a-zA-Z]+\*?", "", t)
    t = t.replace("{", "").replace("}", "")
    t = re.sub(r"\n{3,}", "\n\n", t)
    return t.strip() or "Resume content"


def sanitize_tex(source):
    """Quick sanitize to remove LaTeX "\\" linebreaks followed by a blank line.
    This avoids "There's no line here to end" errors caused by a "\\" followed
    immediately by an empty paragraph.
    """
    if not isinstance(source, str):
        return source
    s = source
    # collapse sequences of a linebreak (\\) followed by an empty line -> keep single newline
    s = re.sub(r"\\\\\s*\n\s*\n", r"\\\\\n", s)
    # if file ends with a trailing \\ then drop the trailing backslash
    s = re.sub(r"\\\\\s*\n\Z", r"\n", s)
    return s


def build_fallback_pdf(text):
    """Minimal pure-Python PDF from plain text (no deps)."""
    max_w = 100
    lines = []
    for raw in str(text or "").splitlines():
        stripped = raw.rstrip()
        if not stripped:
            lines.append("")
        else:
            lines.extend(textwrap.wrap(stripped, width=max_w) or [""])

    PH, TY, LX, LS, BL = 792, 760, 50, 14, 60
    pages, cur, cy = [], [], TY
    for ln in lines:
        if cy < BL:
            pages.append(cur)
            cur = []
            cy = TY
        cur.append((LX, cy, ln))
        cy -= LS
    if cur:
        pages.append(cur)

    def esc(s):
        return s.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")

    objs = []
    objs.append("<< /Type /Catalog /Pages 2 0 R >>")
    pids = [3 + 2 * i for i in range(len(pages))]
    cids = [4 + 2 * i for i in range(len(pages))]
    fid  = 3 + 2 * len(pages)
    kids = " ".join(f"{p} 0 R" for p in pids)
    objs.append(f"<< /Type /Pages /Count {len(pages)} /Kids [{kids}] >>")
    for i, pg in enumerate(pages):
        objs.append(
            f"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 {PH}] "
            f"/Resources << /Font << /F1 {fid} 0 R >> >> /Contents {cids[i]} 0 R >>"
        )
        parts = ["BT", "/F1 10 Tf"]
        for x, y, ln in pg:
            parts.append(f"1 0 0 1 {x} {y} Tm ({esc(ln)}) Tj")
        parts.append("ET")
        s = "\n".join(parts)
        objs.append(f"<< /Length {len(s.encode())} >>\nstream\n{s}\nendstream")
    objs.append("<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>")

    buf = io.BytesIO()
    buf.write(b"%PDF-1.4\n")
    offs = [0]
    for i, obj in enumerate(objs, 1):
        offs.append(buf.tell())
        buf.write(f"{i} 0 obj\n{obj}\nendobj\n".encode())
    xp = buf.tell()
    buf.write(f"xref\n0 {len(objs)+1}\n0000000000 65535 f \n".encode())
    for o in offs[1:]:
        buf.write(f"{o:010d} 00000 n \n".encode())
    buf.write(f"trailer\n<< /Size {len(objs)+1} /Root 1 0 R >>\nstartxref\n{xp}\n%%EOF".encode())
    return buf.getvalue()


def add_hyperlink(paragraph, text, url):
    """Add a clickable hyperlink to a paragraph with 11pt font size."""
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls
    
    part = paragraph.part
    rel_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    
    # Create hyperlink XML with font size
    hyperlink_xml = f'<w:hyperlink {nsdecls("w", "r")} r:id="{rel_id}"><w:r><w:rPr><w:rStyle w:val="Hyperlink"/><w:sz w:val="22"/></w:rPr><w:t>{text}</w:t></w:r></w:hyperlink>'
    
    try:
        hyperlink_element = parse_xml(hyperlink_xml)
        paragraph._element.append(hyperlink_element)
    except Exception:
        # Fallback: just add text if hyperlink creation fails
        run = paragraph.add_run(text)
        run.font.size = Pt(11)


def _set_paragraph_spacing(paragraph, before=0, after=0, line=1.0):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def _add_section_heading(doc, text):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    heading = doc.add_paragraph()
    _set_paragraph_spacing(heading, before=10, after=4, line=1.0)
    
    # Elegant solid bottom border beneath the section heading text
    pPr = heading._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')  # border thickness
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000') # solid black color line
    pbdr.append(bottom)
    pPr.append(pbdr)

    run = heading.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0, 0, 0)


def build_docx_from_latex_data(data):
    doc = Document()
    section = doc.sections[0]
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(10.5)

    name = str(data.get("full_name") or data.get("name") or "YOUR NAME").strip()
    name_p = doc.add_paragraph()
    _set_paragraph_spacing(name_p, before=0, after=2, line=1.0)
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_p.add_run(name.upper())
    name_run.bold = True
    name_run.font.size = Pt(16)

    # Subtitle Line: headline | location
    headline = str(data.get("headline") or "").strip()
    location = str(data.get("location") or "").strip()
    sub_items = []
    if headline:
        sub_items.append(headline)
    if location:
        sub_items.append(location)

    if sub_items:
        sub_p = doc.add_paragraph()
        _set_paragraph_spacing(sub_p, before=0, after=2, line=1.0)
        sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_run = sub_p.add_run(" | ".join(sub_items))
        sub_run.bold = True
        sub_run.font.size = Pt(11)

    # Contact line with no icons, direct clean hyperlinked text
    contact_p = doc.add_paragraph()
    _set_paragraph_spacing(contact_p, before=0, after=8, line=1.0)
    contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    phone = str(data.get("phone") or "").strip()
    email = str(data.get("email") or "").strip()
    linkedin = str(data.get("linkedin") or "").strip()
    github = str(data.get("github") or "").strip()
    
    contact_runs = []
    if phone:
        contact_runs.append(("text", phone))
    if email:
        contact_runs.append(("email", email))
    if linkedin:
        contact_runs.append(("link", ("LinkedIn", linkedin)))
    if github:
        contact_runs.append(("link", ("GitHub", github)))
    
    for idx, item in enumerate(contact_runs):
        if idx > 0:
            contact_p.add_run(" | ")
        
        if item[0] == "text":
            r = contact_p.add_run(item[1])
            r.font.size = Pt(10.5)
        elif item[0] == "email":
            add_hyperlink(contact_p, item[1], f"mailto:{item[1]}")
        elif item[0] == "link":
            label, url = item[1]
            url_clean = url if url.startswith("http") else f"https://{url}"
            add_hyperlink(contact_p, label, url_clean)

    summary = trim_text(data.get("summary") or "", 850)
    if summary:
        _add_section_heading(doc, "Professional Summary")
        p = doc.add_paragraph(summary)
        _set_paragraph_spacing(p, before=0, after=4, line=1.05)
        p.runs[0].font.size = Pt(10.5)

    skills = trim_lines(parse_field_lines(data.get("skills")), max_lines=12, max_len=240)
    if skills:
        _add_section_heading(doc, "Skills")
        for line in skills:
            p = doc.add_paragraph()
            _set_paragraph_spacing(p, before=0, after=2, line=1.0)
            if ":" in line:
                label, value = line.split(":", 1)
                r1 = p.add_run(label.strip() + ": ")
                r1.bold = True
                r1.font.size = Pt(10.5)
                r2 = p.add_run(value.strip())
                r2.font.size = Pt(10.5)
            else:
                r = p.add_run(line)
                r.font.size = Pt(10.5)

    experiences = parse_experience_entries(data.get("experience"))[:8]
    if experiences:
        _add_section_heading(doc, "Professional Experience")
        for exp in experiences:
            title = str(exp.get("title") or "Role").strip()
            company = str(exp.get("company") or "").strip()
            duration = str(exp.get("duration") or "").strip()
            location = str(exp.get("location") or "").strip()

            # Job Title
            tp = doc.add_paragraph()
            _set_paragraph_spacing(tp, before=4, after=1, line=1.0)
            tr = tp.add_run(title)
            tr.bold = True
            tr.font.size = Pt(11.5)

            # Company details with location on same line
            if company:
                cp = doc.add_paragraph()
                _set_paragraph_spacing(cp, before=0, after=1, line=1.0)
                company_text = company
                if location and location != "Location":
                    company_text += f" ({location})"
                cr = cp.add_run(company_text)
                cr.font.size = Pt(11)

            # Employment Dates
            if duration and duration != "Date":
                dp = doc.add_paragraph()
                _set_paragraph_spacing(dp, before=0, after=2, line=1.0)
                dr = dp.add_run(duration)
                dr.italic = True
                dr.font.size = Pt(10)

            # Bullet points
            for bullet in trim_lines(exp.get("bullets") or [], max_lines=8, max_len=260):
                bp = doc.add_paragraph(style="List Bullet")
                _set_paragraph_spacing(bp, before=1, after=1, line=1.1)
                br = bp.add_run(bullet)
                br.font.size = Pt(10.5)

    projects = parse_project_entries(data.get("projects"))[:8]
    if projects:
        _add_section_heading(doc, "Projects")
        for proj in projects:
            name = str(proj.get("name") or "Project").strip()
            tech = str(proj.get("tech") or "").strip()
            date = str(proj.get("date") or "").strip()

            # Name and Date on same line
            pp = doc.add_paragraph()
            _set_paragraph_spacing(pp, before=4, after=1, line=1.0)
            proj_title = name
            if date and date != "Date":
                proj_title += f" — {date}"
            pr = pp.add_run(proj_title)
            pr.bold = True
            pr.font.size = Pt(11.5)

            # Technologies used
            if tech and tech != "Tech Stack":
                tp = doc.add_paragraph()
                _set_paragraph_spacing(tp, before=0, after=2, line=1.0)
                tr = tp.add_run(tech)
                tr.italic = True
                tr.font.size = Pt(10.5)

            # Project details
            for bullet in trim_lines(proj.get("bullets") or [], max_lines=8, max_len=260):
                bp = doc.add_paragraph(style="List Bullet")
                _set_paragraph_spacing(bp, before=1, after=1, line=1.1)
                br = bp.add_run(bullet)
                br.font.size = Pt(10.5)

    education = parse_block_items(data.get("education"), min_fields=3, defaults=["Institution", "Year"])[:6]
    if education:
        _add_section_heading(doc, "Education")
        for row in education:
            degree = str(row[0] if len(row) > 0 else "Degree").strip()
            institution = str(row[1] if len(row) > 1 else "Institution").strip()
            year = str(row[2] if len(row) > 2 else "Year").strip()
            location = str(row[3] if len(row) > 3 else "").strip()

            line = degree
            if institution:
                line += f" | {institution}"
            if year:
                line += f" | {year}"
            if location:
                line += f" | {location}"
            ep = doc.add_paragraph(line)
            _set_paragraph_spacing(ep, before=0, after=1, line=1.0)
            ep.runs[0].font.size = Pt(10.5)

    certs = trim_lines(parse_field_lines(data.get("certifications")), max_lines=14, max_len=230)
    if certs:
        _add_section_heading(doc, "Certifications & Courses")
        for cert in certs:
            bp = doc.add_paragraph(style="List Bullet")
            _set_paragraph_spacing(bp, before=0, after=0, line=1.0)
            br = bp.add_run(cert)
            br.font.size = Pt(10.5)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ── routes ────────────────────────────────────────────────────────────────────

@latex_blueprint.route("/latex-engine-status", methods=["GET"])
def engine_status():
    engines = find_usable_engines()
    return jsonify({
        "available": bool(engines),
        "engines": engines,
        "primary": engines[0] if engines else "",
        "message": f"Ready ({', '.join(engines)})" if engines else "No LaTeX engine found",
    })


@latex_blueprint.route("/export-latex-pdf", methods=["POST"])
def export_pdf():
    data = request.get_json(silent=True) or {}
    latex_data = data.get("latex_data") or {}
    template_id = data.get("template_id") or "classic"

    if not isinstance(latex_data, dict) or not latex_data:
        return jsonify({"error": "latex_data is required"}), 400

    try:
        src = render_template_document(latex_data, template_id=template_id)
    except Exception as exc:
        return jsonify({"error": f"Template render failed: {exc}"}), 400

    if not isinstance(src, str) or not src.strip():
        return jsonify({"error": "Rendered LaTeX source is empty"}), 400
    if len(src) > 300_000:
        return jsonify({"error": "Rendered LaTeX source exceeds 300 KB limit"}), 400

    try:
        src = sanitize_tex(src)
        pdf = compile_latex(src)
        return send_file(
            io.BytesIO(pdf),
            as_attachment=True,
            download_name="resume.pdf",
            mimetype="application/pdf",
        )
    except (FileNotFoundError, subprocess.TimeoutExpired, RuntimeError) as exc:
        fallback = build_fallback_pdf(latex_to_plain(src))
        resp = send_file(
            io.BytesIO(fallback),
            as_attachment=True,
            download_name="resume.pdf",
            mimetype="application/pdf",
        )
        def _sanitize_header(s, maxlen=300):
            s = str(s or "")
            s = s.replace("\r", " ").replace("\n", " ")
            s = re.sub(r"\s+", " ", s).strip()
            if len(s) > maxlen:
                s = s[:maxlen-3] + "..."
            return s
        resp.headers["X-Latex-Fallback"] = "plain-text"
        resp.headers["X-Latex-Error"] = _sanitize_header(str(exc))
        return resp
    except Exception as exc:
        return jsonify({"error": f"Unexpected error: {exc}"}), 500


@latex_blueprint.route("/export-latex-docx", methods=["POST"])
def export_latex_docx():
    data = request.get_json(silent=True) or {}
    latex_data = data.get("latex_data") or {}

    if not isinstance(latex_data, dict) or not latex_data:
        return jsonify({"error": "latex_data is required"}), 400

    try:
        docx_buf = build_docx_from_latex_data(latex_data)
        return send_file(
            docx_buf,
            as_attachment=True,
            download_name="resume.docx",
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    except Exception as exc:
        return jsonify({"error": f"DOCX generation failed: {exc}"}), 500


@latex_blueprint.route("/validate-latex", methods=["POST"])
def validate_latex():
    """Dry-run compile — returns errors without sending a PDF."""
    data = request.get_json(silent=True) or {}
    latex_data = data.get("latex_data") or {}
    template_id = data.get("template_id") or "classic"
    if not isinstance(latex_data, dict) or not latex_data:
        return jsonify({"valid": False, "error": "latex_data is required"}), 400
    try:
        src = render_template_document(latex_data, template_id=template_id)
    except Exception as exc:
        return jsonify({"valid": False, "error": f"Template render failed: {exc}"}), 400
    if not isinstance(src, str) or not src.strip():
        return jsonify({"valid": False, "error": "No source provided"}), 400

    engines = find_usable_engines()
    if not engines:
        return jsonify({"valid": False, "error": "No LaTeX engine available"}), 503

    src = sanitize_tex(src)
    pdf, log = compile_once(engines[0], src)
    if pdf:
        return jsonify({"valid": True, "engine": engines[0]})
    missing = extract_missing_packages(log)
    return jsonify({
        "valid": False,
        "engine": engines[0],
        "missing_packages": missing,
        "log_tail": last_lines(log, 20),
    }), 422
